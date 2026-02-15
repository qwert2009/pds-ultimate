"""
PDS-Ultimate Business Tools
==============================
Регистрация бизнес-инструментов для AI-агента.

Каждый модуль системы (заказы, финансы, логистика, секретарь)
экспортирует свои возможности как формальные Tool-ы.

Агент (ReAct loop) вызывает их через ToolRegistry.
Это обеспечивает:
- Формальный контракт (параметры, описание)
- Единую точку входа для LLM
- Логирование и обработку ошибок
- Масштабируемость (новые tools = новые возможности)
"""

from __future__ import annotations

import asyncio
import os
from datetime import date, timedelta

from pds_ultimate.config import config, logger
from pds_ultimate.core.tools import Tool, ToolParameter, ToolResult, tool_registry

# ═══════════════════════════════════════════════════════════════════════════════
# ЛОГИСТИКА / ЗАКАЗЫ
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_create_order(items_text: str, db_session=None) -> ToolResult:
    """Создать новый заказ из текстового описания позиций."""
    from pds_ultimate.core.database import (
        ItemStatus,
        Order,
        OrderItem,
        OrderStatus,
    )
    from pds_ultimate.utils.parsers import parser

    if not db_session:
        return ToolResult("create_order", False, "", error="Нет сессии БД")

    result = await parser.parse_text_smart(items_text)
    if not result.items:
        # Пробуем через LLM
        from pds_ultimate.core.llm_engine import llm_engine
        parsed = await llm_engine.parse_order(items_text)
        if not parsed:
            return ToolResult("create_order", False, "",
                              error="Не удалось распознать позиции")
        items_data = parsed
    else:
        items_data = [item.to_dict() for item in result.items]

    order_count = db_session.query(Order).count()
    order_number = f"ORD-{order_count + 1:04d}"

    order = Order(
        order_number=order_number,
        status=OrderStatus.CONFIRMED,
        order_date=date.today(),
    )
    db_session.add(order)
    db_session.flush()

    created_items = []
    for item_data in items_data:
        first_check = date.today() + timedelta(days=config.logistics.first_status_check_days)
        item = OrderItem(
            order_id=order.id,
            name=item_data.get("name", item_data.get("name", "?")),
            quantity=float(item_data.get("quantity", 1)),
            unit=item_data.get("unit", "шт"),
            unit_price=item_data.get("unit_price"),
            price_currency=item_data.get("currency", "USD"),
            weight=item_data.get("weight"),
            status=ItemStatus.PENDING,
            next_check_date=first_check,
        )
        db_session.add(item)
        created_items.append(item_data)

    db_session.commit()

    items_text_lines = "\n".join(
        f"  {i + 1}. {it.get('name', '?')} — {it.get('quantity', '?')} {it.get('unit', 'шт')}"
        for i, it in enumerate(created_items)
    )

    return ToolResult(
        "create_order",
        True,
        f"✅ Заказ {order_number} создан ({len(created_items)} позиций):\n{items_text_lines}",
        data={"order_id": order.id, "order_number": order_number,
              "items_count": len(created_items)},
    )


async def tool_get_orders_status(order_number: str = None, db_session=None) -> ToolResult:
    """Получить статус заказов."""
    from pds_ultimate.core.database import (
        ItemStatus,
        Order,
        OrderItem,
        OrderStatus,
    )

    if not db_session:
        return ToolResult("get_orders_status", False, "", error="Нет сессии БД")

    if order_number:
        order = db_session.query(Order).filter_by(
            order_number=order_number).first()
        if not order:
            return ToolResult("get_orders_status", False, "",
                              error=f"Заказ {order_number} не найден")

        items = db_session.query(OrderItem).filter_by(order_id=order.id).all()
        items_info = []
        for item in items:
            emoji = "✅" if item.status == ItemStatus.ARRIVED else "⏳"
            track = f" | Трек: {item.tracking_number}" if item.tracking_number else ""
            items_info.append(
                f"  {emoji} {item.name} — {item.quantity} {item.unit}{track}")

        text = (
            f"📦 Заказ {order.order_number}\n"
            f"Статус: {order.status.value}\n"
            f"Дата: {order.order_date}\n"
            f"Позиции:\n" + "\n".join(items_info)
        )
        if order.income:
            text += f"\n💰 Доход: {order.income} {order.income_currency}"
        if order.net_profit is not None:
            text += f"\n📊 Чистая прибыль: ${order.net_profit:.2f}"

        return ToolResult("get_orders_status", True, text,
                          data={"order": order.order_number, "status": order.status.value})

    # Все активные
    active = db_session.query(Order).filter(
        Order.status.notin_([OrderStatus.ARCHIVED, OrderStatus.COMPLETED])
    ).all()

    if not active:
        return ToolResult("get_orders_status", True, "Нет активных заказов.")

    lines = ["📋 Активные заказы:\n"]
    for o in active:
        item_count = db_session.query(
            OrderItem).filter_by(order_id=o.id).count()
        pending = db_session.query(OrderItem).filter_by(
            order_id=o.id, status=ItemStatus.PENDING).count()
        lines.append(
            f"• {o.order_number} | {o.status.value} | Позиций: {item_count} (ждём: {pending})")

    return ToolResult("get_orders_status", True, "\n".join(lines),
                      data={"active_count": len(active)})


async def tool_set_income(order_number: str, amount: float,
                          currency: str = "USD", db_session=None) -> ToolResult:
    """Установить доход за заказ."""
    from pds_ultimate.core.database import Order, Transaction, TransactionType

    if not db_session:
        return ToolResult("set_income", False, "", error="Нет сессии БД")

    order = db_session.query(Order).filter_by(
        order_number=order_number).first()
    if not order:
        return ToolResult("set_income", False, "",
                          error=f"Заказ {order_number} не найден")

    order.income = amount
    order.income_currency = currency

    amount_usd = _convert_to_usd(amount, currency)
    db_session.add(Transaction(
        order_id=order.id,
        transaction_type=TransactionType.INCOME,
        amount=amount,
        currency=currency,
        amount_usd=amount_usd,
        description=f"Оплата за заказ {order.order_number}",
        transaction_date=date.today(),
    ))
    db_session.commit()

    return ToolResult("set_income", True,
                      f"✅ Доход за {order_number}: {amount} {currency} (${amount_usd:.2f})",
                      data={"order": order_number, "amount_usd": amount_usd})


async def tool_set_expense(order_number: str, amount: float,
                           currency: str = "USD", db_session=None) -> ToolResult:
    """Установить расход на товар."""
    from pds_ultimate.core.database import (
        Order,
        OrderStatus,
        Transaction,
        TransactionType,
    )

    if not db_session:
        return ToolResult("set_expense", False, "", error="Нет сессии БД")

    order = db_session.query(Order).filter_by(
        order_number=order_number).first()
    if not order:
        return ToolResult("set_expense", False, "",
                          error=f"Заказ {order_number} не найден")

    order.expense_goods = amount
    order.expense_goods_currency = currency

    amount_usd = _convert_to_usd(amount, currency)
    db_session.add(Transaction(
        order_id=order.id,
        transaction_type=TransactionType.EXPENSE_GOODS,
        amount=amount,
        currency=currency,
        amount_usd=amount_usd,
        description=f"Оплата поставщику за {order.order_number}",
        transaction_date=date.today(),
    ))

    income_usd = _convert_to_usd(
        order.income or 0, order.income_currency or "USD")
    remainder = income_usd - amount_usd

    order.status = OrderStatus.TRACKING
    db_session.commit()

    return ToolResult("set_expense", True,
                      f"✅ Расход на товар: {amount} {currency}\n📊 Остаток: ${remainder:.2f}",
                      data={"order": order_number, "remainder_usd": remainder})


# ═══════════════════════════════════════════════════════════════════════════════
# ФИНАНСЫ
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_get_financial_summary(db_session=None) -> ToolResult:
    """Получить финансовую сводку."""
    from sqlalchemy import func

    from pds_ultimate.core.database import (
        Order,
        OrderStatus,
        Transaction,
        TransactionType,
    )

    if not db_session:
        return ToolResult("get_financial_summary", False, "", error="Нет сессии БД")

    total_income = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.INCOME).scalar() or 0

    total_goods = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.EXPENSE_GOODS).scalar() or 0

    total_delivery = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.EXPENSE_DELIVERY).scalar() or 0

    total_savings = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.PROFIT_SAVINGS).scalar() or 0

    total_profit_exp = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.PROFIT_EXPENSES).scalar() or 0

    completed = db_session.query(Order).filter(
        Order.status.in_([OrderStatus.COMPLETED, OrderStatus.ARCHIVED])
    ).count()

    active = db_session.query(Order).filter(
        Order.status.notin_([OrderStatus.ARCHIVED, OrderStatus.COMPLETED])
    ).count()

    net = total_income - total_goods - total_delivery

    text = (
        f"💰 ФИНАНСОВАЯ СВОДКА (USD)\n\n"
        f"Общий доход: ${total_income:.2f}\n"
        f"Расходы на товар: ${total_goods:.2f}\n"
        f"Расходы на доставку: ${total_delivery:.2f}\n"
        f"━━━━━━━━━━━━━━━\n"
        f"Чистая прибыль: ${net:.2f}\n\n"
        f"На расходы: ${total_profit_exp:.2f}\n"
        f"Отложено: ${total_savings:.2f}\n\n"
        f"Активных заказов: {active}\n"
        f"Закрытых: {completed}"
    )

    return ToolResult("get_financial_summary", True, text, data={
        "income": total_income, "goods": total_goods,
        "delivery": total_delivery, "net_profit": net,
        "savings": total_savings, "active_orders": active,
    })


async def tool_convert_currency(amount: float, from_currency: str,
                                to_currency: str = "USD", **kwargs) -> ToolResult:
    """Конвертировать валюту."""
    rates = {"TMT": 19.5, "CNY": 7.1}

    # from → USD
    if from_currency == "USD":
        usd = amount
    elif from_currency in rates:
        usd = amount / rates[from_currency]
    else:
        return ToolResult("convert_currency", False, "",
                          error=f"Неизвестная валюта: {from_currency}")

    # USD → to
    if to_currency == "USD":
        result_amount = usd
    elif to_currency in rates:
        result_amount = usd * rates[to_currency]
    else:
        return ToolResult("convert_currency", False, "",
                          error=f"Неизвестная валюта: {to_currency}")

    return ToolResult("convert_currency", True,
                      f"{amount} {from_currency} = {result_amount:.2f} {to_currency}",
                      data={"result": result_amount, "currency": to_currency})


# ═══════════════════════════════════════════════════════════════════════════════
# КОНТАКТЫ
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_save_contact_note(name: str, note: str, is_warning: bool = False,
                                 db_session=None) -> ToolResult:
    """Сохранить заметку о контакте."""
    from pds_ultimate.core.database import Contact, ContactType

    if not db_session:
        return ToolResult("save_contact_note", False, "", error="Нет сессии БД")

    contact = db_session.query(Contact).filter(
        Contact.name.ilike(f"%{name}%")
    ).first()

    if not contact:
        contact = Contact(name=name, contact_type=ContactType.OTHER)
        db_session.add(contact)
        db_session.flush()

    today = date.today()
    if is_warning:
        existing = contact.warnings or ""
        contact.warnings = f"{existing}\n[{today}] {note}".strip()
    else:
        existing = contact.notes or ""
        contact.notes = f"{existing}\n[{today}] {note}".strip()

    db_session.commit()

    emoji = "⚠️" if is_warning else "📝"
    return ToolResult("save_contact_note", True,
                      f"{emoji} Записал о «{contact.name}»: {note}")


async def tool_find_contact(query: str, db_session=None) -> ToolResult:
    """Найти контакт по имени."""
    from pds_ultimate.core.database import Contact

    if not db_session:
        return ToolResult("find_contact", False, "", error="Нет сессии БД")

    contacts = db_session.query(Contact).filter(
        Contact.name.ilike(f"%{query}%")
    ).limit(10).all()

    if not contacts:
        return ToolResult("find_contact", True, f"Контакт «{query}» не найден.")

    lines = [f"🔍 Найдено ({len(contacts)}):"]
    for c in contacts:
        info = f"• {c.name} ({c.contact_type.value})"
        if c.phone:
            info += f" | {c.phone}"
        if c.warnings:
            info += " ⚠️"
        if c.notes:
            last_note = c.notes.strip().split("\n")[-1]
            info += f"\n  📝 {last_note[:80]}"
        lines.append(info)

    return ToolResult("find_contact", True, "\n".join(lines),
                      data={"count": len(contacts)})


# ═══════════════════════════════════════════════════════════════════════════════
# КАЛЕНДАРЬ & НАПОМИНАНИЯ
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_create_reminder(message: str, scheduled_at: str,
                               db_session=None) -> ToolResult:
    """Создать напоминание."""
    from datetime import datetime

    from pds_ultimate.core.database import Reminder, ReminderStatus

    if not db_session:
        return ToolResult("create_reminder", False, "", error="Нет сессии БД")

    try:
        # Пробуем разные форматы даты
        dt = None
        for fmt in ("%Y-%m-%d %H:%M", "%d.%m.%Y %H:%M", "%Y-%m-%dT%H:%M", "%d.%m.%Y"):
            try:
                dt = datetime.strptime(scheduled_at, fmt)
                break
            except ValueError:
                continue

        if not dt:
            return ToolResult("create_reminder", False, "",
                              error=f"Не распознан формат даты: {scheduled_at}")

        reminder = Reminder(
            message=message,
            scheduled_at=dt,
            status=ReminderStatus.PENDING,
            reminder_minutes=30,
        )
        db_session.add(reminder)
        db_session.commit()

        return ToolResult("create_reminder", True,
                          f"⏰ Напоминание создано: «{message}» на {dt.strftime('%d.%m.%Y %H:%M')}",
                          data={"reminder_id": reminder.id})

    except Exception as e:
        return ToolResult("create_reminder", False, "", error=str(e))


async def tool_create_calendar_event(title: str, event_date: str,
                                     description: str = "",
                                     db_session=None) -> ToolResult:
    """Создать событие в календаре."""
    from datetime import datetime

    from pds_ultimate.core.database import CalendarEvent

    if not db_session:
        return ToolResult("create_calendar_event", False, "", error="Нет сессии БД")

    try:
        dt = None
        for fmt in ("%Y-%m-%d %H:%M", "%d.%m.%Y %H:%M", "%Y-%m-%dT%H:%M", "%d.%m.%Y"):
            try:
                dt = datetime.strptime(event_date, fmt)
                break
            except ValueError:
                continue

        if not dt:
            return ToolResult("create_calendar_event", False, "",
                              error=f"Не распознан формат даты: {event_date}")

        event = CalendarEvent(
            title=title,
            event_date=dt,
            description=description,
            reminder_minutes=30,
        )
        db_session.add(event)
        db_session.commit()

        return ToolResult("create_calendar_event", True,
                          f"📅 Событие создано: «{title}» на {dt.strftime('%d.%m.%Y %H:%M')}",
                          data={"event_id": event.id})

    except Exception as e:
        return ToolResult("create_calendar_event", False, "", error=str(e))


# ═══════════════════════════════════════════════════════════════════════════════
# УТРЕННИЙ БРИФИНГ & ОТЧЁТЫ
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_morning_brief(db_session=None) -> ToolResult:
    """Сформировать утренний брифинг."""
    from sqlalchemy import func

    from pds_ultimate.core.database import (
        ItemStatus,
        Order,
        OrderItem,
        OrderStatus,
        Transaction,
        TransactionType,
    )

    if not db_session:
        return ToolResult("morning_brief", False, "", error="Нет сессии БД")

    active_orders = db_session.query(Order).filter(
        Order.status.notin_([OrderStatus.ARCHIVED, OrderStatus.COMPLETED])
    ).count()

    pending_items = db_session.query(OrderItem).filter_by(
        status=ItemStatus.PENDING
    ).count()

    total_income = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.INCOME).scalar() or 0

    total_expenses = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter(Transaction.transaction_type.in_([
        TransactionType.EXPENSE_GOODS,
        TransactionType.EXPENSE_DELIVERY,
    ])).scalar() or 0

    total_savings = db_session.query(
        func.sum(Transaction.amount_usd)
    ).filter_by(transaction_type=TransactionType.PROFIT_SAVINGS).scalar() or 0

    balance = total_income - total_expenses
    today = date.today().strftime("%d.%m.%Y")

    text = (
        f"☀️ БРИФИНГ НА {today}\n\n"
        f"📦 Активных заказов: {active_orders}\n"
        f"📋 Ожидаем позиций: {pending_items}\n"
        f"💰 Баланс: ${balance:.2f}\n"
        f"🏦 Отложено: ${total_savings:.2f}\n\n"
        f"Что делаем сегодня, босс?"
    )

    return ToolResult("morning_brief", True, text, data={
        "active_orders": active_orders, "pending_items": pending_items,
        "balance": balance, "savings": total_savings,
    })


# ═══════════════════════════════════════════════════════════════════════════════
# ПЕРЕВОД & ТЕКСТ
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_translate(text: str, target_lang: str = "ru",
                         source_lang: str = None, **kwargs) -> ToolResult:
    """Перевести текст."""
    from pds_ultimate.core.llm_engine import llm_engine

    result = await llm_engine.translate(text, target_lang, source_lang)
    return ToolResult("translate", True, result,
                      data={"target_lang": target_lang})


async def tool_summarize(text: str, **kwargs) -> ToolResult:
    """Создать краткое саммари текста."""
    from pds_ultimate.core.llm_engine import llm_engine

    result = await llm_engine.summarize(text)
    return ToolResult("summarize", True, result)


# ═══════════════════════════════════════════════════════════════════════════════
# БЕЗОПАСНОСТЬ
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_security_emergency(db_session=None) -> ToolResult:
    """Активировать экстренный режим безопасности."""

    from pds_ultimate.config import ALL_ORDERS_ARCHIVE_PATH, MASTER_FINANCE_PATH
    from pds_ultimate.core.database import Transaction

    if not db_session:
        return ToolResult("security_emergency", False, "", error="Нет сессии БД")

    for fp in [MASTER_FINANCE_PATH, ALL_ORDERS_ARCHIVE_PATH]:
        if fp.exists():
            try:
                os.remove(fp)
            except OSError:
                pass

    db_session.query(Transaction).delete()
    db_session.commit()

    logger.critical("🚨 SECURITY MODE ACTIVATED")
    return ToolResult("security_emergency", True,
                      "🔒 Режим безопасности активирован. Финансовые данные удалены.")


# ═══════════════════════════════════════════════════════════════════════════════
# ПАМЯТЬ АГЕНТА (tools для работы с долгосрочной памятью)
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_remember(fact: str, importance: float = 0.5,
                        memory_type: str = "fact", **kwargs) -> ToolResult:
    """Запомнить важный факт."""
    from pds_ultimate.core.memory import memory_manager

    entry = memory_manager.store_fact(
        content=fact,
        importance=importance,
        tags=[memory_type],
        source="agent",
    )
    return ToolResult("remember", True,
                      f"📌 Запомнил: «{fact}» (важность: {importance})")


async def tool_recall(query: str, **kwargs) -> ToolResult:
    """Вспомнить факты по запросу."""
    from pds_ultimate.core.memory import memory_manager

    entries = memory_manager.recall(query, limit=5)
    if not entries:
        return ToolResult("recall", True, "Ничего не найдено в памяти.")

    lines = ["🧠 Вспомнил:"]
    for e in entries:
        lines.append(f"  • [{e.memory_type}] {e.content}")

    return ToolResult("recall", True, "\n".join(lines),
                      data=[e.to_dict() for e in entries])


# ═══════════════════════════════════════════════════════════════════════════════
# УТИЛИТЫ
# ═══════════════════════════════════════════════════════════════════════════════

def _convert_to_usd(amount: float, currency: str) -> float:
    """Конвертировать в USD."""
    if currency == "USD":
        return amount
    rates = config.currency.fixed_rates
    if currency in rates:
        return round(amount / rates[currency], 2)
    return amount


# ═══════════════════════════════════════════════════════════════════════════════
# PART 7: NEW TOOL HANDLERS
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_exchange_rates(
    from_currency: str = "USD",
    to_currency: str = "",
    amount: float = 1.0,
    **kwargs,
) -> ToolResult:
    """Получить курс обмена валют (онлайн + кэш + фиксированные)."""
    from pds_ultimate.integrations.exchange_rates import exchange_service

    try:
        if to_currency:
            result = await exchange_service.convert(
                amount, from_currency.upper(), to_currency.upper()
            )
            if "error" in result:
                return ToolResult(
                    "exchange_rates", False, "",
                    error=result["error"],
                )
            return ToolResult(
                "exchange_rates", True,
                f"💱 {amount:.2f} {from_currency.upper()} = "
                f"{result['result']:.2f} {to_currency.upper()}\n"
                f"Курс: {result['rate']:.4f} "
                f"(источник: {result.get('source', 'unknown')})",
                data=result,
            )

        result = await exchange_service.refresh_all()
        table = exchange_service.format_rates_table()
        return ToolResult(
            "exchange_rates", True, table,
            data={"rates_count": len(result.rates)},
        )

    except Exception as e:
        return ToolResult(
            "exchange_rates", False, "",
            error=f"Ошибка получения курсов: {e}",
        )


async def tool_ocr_recognize(
    file_path: str,
    extract_amounts: bool = False,
    extract_tracking: bool = False,
    **kwargs,
) -> ToolResult:
    """Распознать текст на изображении (OCR)."""
    from pds_ultimate.modules.files.ocr_engine import ocr_engine

    try:
        result = await ocr_engine.recognize(file_path)
        lines = [f"📝 OCR ({result.engine_used})"]
        lines.append(f"Уверенность: {result.avg_confidence:.0%}")
        lines.append(f"\n{result.confident_text[:2000]}")

        data = {"text": result.confident_text,
                "confidence": result.avg_confidence}

        if extract_amounts:
            amounts = await ocr_engine.extract_amounts(file_path)
            if amounts:
                lines.append("\n💰 Суммы:")
                for a in amounts:
                    lines.append(f"  {a.original} → {a.amount} {a.currency}")
                data["amounts"] = [
                    {"amount": a.amount, "currency": a.currency}
                    for a in amounts
                ]

        if extract_tracking:
            tracking = await ocr_engine.extract_tracking_numbers(file_path)
            if tracking:
                lines.append("\n📦 Трекинг:")
                for t in tracking:
                    lines.append(f"  {t.number} ({t.carrier})")
                data["tracking"] = [
                    {"number": t.number, "carrier": t.carrier}
                    for t in tracking
                ]

        return ToolResult(
            "ocr_recognize", True, "\n".join(lines), data=data,
        )

    except Exception as e:
        return ToolResult(
            "ocr_recognize", False, "",
            error=f"Ошибка OCR: {e}",
        )


async def tool_scan_receipt(
    file_path: str,
    save_to_db: bool = True,
    db_session=None,
    **kwargs,
) -> ToolResult:
    """Сканировать чек и распознать расходы."""
    from pds_ultimate.modules.executive.receipt_scanner import receipt_scanner

    try:
        receipt = await receipt_scanner.scan_receipt(file_path)
        text = receipt_scanner.format_receipt(receipt)

        if save_to_db and db_session and receipt.amount:
            saved = await receipt_scanner.save_expense(
                receipt, db_session
            )
            if saved:
                text += "\n\n💾 Сохранено в базу расходов"

        return ToolResult(
            "scan_receipt", True, text,
            data={
                "amount": receipt.amount,
                "currency": receipt.currency,
                "category": receipt.category.value if receipt.category else None,
                "vendor": receipt.vendor,
            },
        )

    except Exception as e:
        return ToolResult(
            "scan_receipt", False, "",
            error=f"Ошибка сканирования чека: {e}",
        )


async def tool_translate_text(
    text: str,
    target_lang: str = "ru",
    source_lang: str = "",
    **kwargs,
) -> ToolResult:
    """Перевести текст через TranslatorService (с бизнес-глоссарием)."""
    from pds_ultimate.modules.executive.translator import translator

    try:
        result = await translator.translate(
            text, target_lang, source_lang or None,
        )
        formatted = translator.format_translation(result)
        return ToolResult(
            "translate_text", True, formatted,
            data={
                "source_lang": result.source_lang,
                "target_lang": result.target_lang,
                "translated": result.translated,
            },
        )

    except Exception as e:
        return ToolResult(
            "translate_text", False, "",
            error=f"Ошибка перевода: {e}",
        )


async def tool_archivist_rename(
    file_path: str,
    description: str = "",
    **kwargs,
) -> ToolResult:
    """Стандартизировать имя файла по корпоративному стандарту."""
    from pds_ultimate.modules.executive.archivist import archivist

    try:
        result = archivist.rename_file(file_path, context=description)
        text = archivist.format_rename_result(result)

        if not result.success:
            return ToolResult(
                "archivist_rename", False, text,
                data=result.to_dict(),
                error=result.error or "Не удалось переименовать",
            )

        return ToolResult(
            "archivist_rename", True, text,
            data=result.to_dict(),
        )

    except Exception as e:
        return ToolResult(
            "archivist_rename", False, "",
            error=f"Ошибка переименования: {e}",
        )


async def tool_convert_file(
    file_path: str,
    target_format: str,
    **kwargs,
) -> ToolResult:
    """Конвертировать файл в другой формат."""
    from pds_ultimate.modules.files.converter import file_converter

    try:
        result = await file_converter.convert(file_path, target_format)
        text = file_converter.format_result(result)

        if result.success:
            return ToolResult(
                "convert_file", True, text,
                data=result.to_dict(),
            )
        return ToolResult(
            "convert_file", False, "",
            error=text,
        )

    except Exception as e:
        return ToolResult(
            "convert_file", False, "",
            error=f"Ошибка конвертации: {e}",
        )


async def tool_google_calendar_events(
    action: str = "today",
    title: str = "",
    start_time: str = "",
    end_time: str = "",
    description: str = "",
    **kwargs,
) -> ToolResult:
    """Работа с Google Calendar (создать/просмотреть события)."""
    from pds_ultimate.integrations.google_calendar import google_calendar

    try:
        if action == "today":
            events = await google_calendar.get_today_events()
            text = google_calendar.format_day_summary(events)
            return ToolResult(
                "google_calendar", True, text,
                data={"events_count": len(events)},
            )

        elif action == "create":
            from datetime import datetime

            if not title or not start_time:
                return ToolResult(
                    "google_calendar", False, "",
                    error="Для создания нужны title и start_time",
                )

            # Parse dates
            from pds_ultimate.utils.validators import parse_date
            start_dt = parse_date(start_time)
            end_dt = parse_date(end_time) if end_time else None
            if not start_dt:
                return ToolResult(
                    "google_calendar", False, "",
                    error=f"Не распознан формат даты: {start_time}",
                )

            created = await google_calendar.create_event(
                summary=title,
                start=start_dt,
                end=end_dt,
                description=description,
            )
            if created:
                return ToolResult(
                    "google_calendar", True,
                    f"📅 Событие создано: «{title}»",
                    data={"event_id": created.id},
                )
            return ToolResult(
                "google_calendar", False, "",
                error="Не удалось создать событие",
            )

        elif action == "free_slots":
            from datetime import datetime

            from pds_ultimate.utils.validators import parse_date
            dt = parse_date(start_time) if start_time else datetime.now()
            ref_date = dt or datetime.now()

            # Get today's events first, then find free slots (sync method)
            events = await google_calendar.get_events(
                ref_date.replace(hour=0, minute=0, second=0, microsecond=0),
            )
            slots = google_calendar.find_free_slots(
                events, reference_date=ref_date,
            )
            if slots:
                text = google_calendar.format_free_slots(slots)
                return ToolResult(
                    "google_calendar", True, text,
                    data={"slots_count": len(slots)},
                )
            return ToolResult(
                "google_calendar", True, "Нет свободных слотов на эту дату.",
            )

        return ToolResult(
            "google_calendar", False, "",
            error=f"Неизвестное действие: {action}",
        )

    except Exception as e:
        return ToolResult(
            "google_calendar", False, "",
            error=f"Ошибка Google Calendar: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# MESSAGING & FILES — TOOLS (ОТПРАВКА СООБЩЕНИЙ, СОЗДАНИЕ ФАЙЛОВ, EMAIL)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_send_whatsapp(
    contact_name: str = "",
    phone: str = "",
    message: str = "",
    db_session=None,
    **kwargs,
) -> ToolResult:
    """Отправить сообщение в WhatsApp через Green-API."""
    from pds_ultimate.integrations.whatsapp import wa_client

    if not message:
        return ToolResult("send_whatsapp", False, "", error="Нужен текст сообщения")

    # Определяем chat_id
    chat_id = ""
    if phone:
        # Убираем +, пробелы
        clean = phone.replace("+", "").replace(" ", "").replace("-", "")
        chat_id = f"{clean}@c.us"
    elif contact_name and db_session:
        # Ищем контакт в БД
        from pds_ultimate.core.database import Contact
        contact = db_session.query(Contact).filter(
            Contact.name.ilike(f"%{contact_name}%")
        ).first()
        if contact and contact.whatsapp_id:
            chat_id = contact.whatsapp_id
        elif contact and contact.phone:
            clean = contact.phone.replace(
                "+", "").replace(" ", "").replace("-", "")
            chat_id = f"{clean}@c.us"
        else:
            return ToolResult(
                "send_whatsapp", False, "",
                error=f"Контакт '{contact_name}' не найден или не имеет номера WhatsApp. "
                f"Укажи phone (номер телефона) явно.",
            )
    else:
        return ToolResult(
            "send_whatsapp", False, "",
            error="Укажи contact_name (имя контакта) или phone (номер телефона)",
        )

    if not wa_client._started:
        try:
            await wa_client.start()
        except Exception as e:
            return ToolResult("send_whatsapp", False, "", error=f"WhatsApp не подключён: {e}")

    success = await wa_client.send_message(chat_id, message)
    if success:
        return ToolResult(
            "send_whatsapp", True,
            f"✅ Сообщение отправлено в WhatsApp ({chat_id}):\n«{message[:200]}»",
        )
    return ToolResult("send_whatsapp", False, "", error="Не удалось отправить. Проверь авторизацию Green-API.")


async def tool_read_telegram_chat(
    username: str = "",
    chat_id: int = 0,
    contact_name: str = "",
    limit: int = 20,
    days: int = 30,
    db_session=None,
    **kwargs,
) -> ToolResult:
    """
    Прочитать историю чата Telegram через Telethon userbot.
    Работает по username (@milana), chat_id или contact_name (имя).
    """
    try:
        from pds_ultimate.integrations.telethon_client import telethon_client

        if not telethon_client._started:
            return ToolResult(
                "read_telegram_chat", False, "",
                error="Telethon userbot не запущен — чтение чатов невозможно.",
            )

        # Smart resolve by name
        if not username and not chat_id and contact_name:
            from pds_ultimate.core.contact_book import contact_book
            resolved = contact_book.resolve(
                contact_name, db_session=db_session)
            if resolved:
                username = resolved.get("telegram", "")
                chat_id = resolved.get("telegram_id", 0)

        identifier = chat_id or username
        if not identifier:
            return ToolResult(
                "read_telegram_chat", False, "",
                error="Нужен username или chat_id.",
            )

        messages = await telethon_client.get_messages(
            str(identifier), limit=limit, offset_days=days,
        )

        if not messages:
            return ToolResult(
                "read_telegram_chat", True,
                f"💬 Чат с {username or chat_id}: сообщений за {days} дн. не найдено.",
            )

        lines = [f"💬 Чат с {username or chat_id} (последние {len(messages)}):"]
        for m in messages[:limit]:
            who = "🔵 Я" if m.get("is_owner") else "⚪ Собеседник"
            date_str = m.get("date", "")[:16].replace("T", " ")
            text_preview = (m.get("text") or "")[:200]
            lines.append(f"  {who} [{date_str}]: {text_preview}")

        return ToolResult(
            "read_telegram_chat", True,
            "\n".join(lines),
            data={"messages": messages[:limit]},
        )
    except Exception as e:
        return ToolResult(
            "read_telegram_chat", False, "",
            error=f"Ошибка чтения чата: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# КОНТАКТНАЯ КНИГА (Smart Name → Contact Resolution)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_link_contact(
    name: str = "",
    telegram: str = "",
    phone: str = "",
    email: str = "",
    whatsapp: str = "",
    db_session=None,
    **kwargs,
) -> ToolResult:
    """
    Привязать контактные данные к имени.
    'запомни что у Миланы телеграм @milana_sagomonyan'
    'сохрани что email Кирилла — kirill@mail.ru'
    """
    if not name:
        return ToolResult("link_contact", False, "", error="Нужно указать имя контакта.")

    if not any([telegram, phone, email, whatsapp]):
        return ToolResult(
            "link_contact", False, "",
            error="Нужно указать хотя бы один контакт: telegram, phone, email или whatsapp.",
        )

    from pds_ultimate.core.contact_book import contact_book

    result = contact_book.link(
        name=name,
        telegram=telegram,
        phone=phone,
        email=email,
        whatsapp=whatsapp,
        db_session=db_session,
    )

    if "error" in result:
        return ToolResult("link_contact", False, "", error=result["error"])

    parts = [f"✅ Контакт «{result.get('name', name)}» обновлён:"]
    if result.get("telegram"):
        parts.append(f"  📱 Telegram: @{result['telegram']}")
    if result.get("phone"):
        parts.append(f"  📞 Телефон: {result['phone']}")
    if result.get("email"):
        parts.append(f"  📧 Email: {result['email']}")
    if result.get("whatsapp"):
        parts.append(f"  💬 WhatsApp: {result['whatsapp']}")

    return ToolResult("link_contact", True, "\n".join(parts))


async def tool_resolve_contact(
    name: str = "",
    db_session=None,
    **kwargs,
) -> ToolResult:
    """
    Найти контакт по имени, нику или прозвищу.
    Поддерживает падежи и уменьшительные: 'Милане', 'Серёга', 'Кирюха'.
    """
    if not name:
        return ToolResult("resolve_contact", False, "", error="Нужно указать имя.")

    from pds_ultimate.core.contact_book import contact_book

    contact = contact_book.resolve(name, db_session=db_session)
    if not contact:
        return ToolResult(
            "resolve_contact", True,
            f"🔍 Контакт «{name}» не найден. "
            f"Привяжи: 'запомни что у {name} телеграм @username'",
        )

    parts = [f"📇 {contact.get('name', name)}:"]
    if contact.get("telegram"):
        parts.append(f"  📱 Telegram: @{contact['telegram']}")
    if contact.get("telegram_id"):
        parts.append(f"  🆔 TG ID: {contact['telegram_id']}")
    if contact.get("phone"):
        parts.append(f"  📞 Телефон: {contact['phone']}")
    if contact.get("email"):
        parts.append(f"  📧 Email: {contact['email']}")
    if contact.get("whatsapp"):
        parts.append(f"  💬 WhatsApp: {contact['whatsapp']}")
    if contact.get("notes"):
        parts.append(f"  📝 {contact['notes'][:100]}")

    return ToolResult("resolve_contact", True, "\n".join(parts), data=contact)


async def tool_list_contacts(db_session=None, **kwargs) -> ToolResult:
    """Показать все контакты из адресной книги."""
    from pds_ultimate.core.contact_book import contact_book

    contacts = contact_book.list_all()
    if not contacts:
        return ToolResult("list_contacts", True, "📇 Адресная книга пуста.")

    lines = [f"📇 Адресная книга ({len(contacts)} контактов):"]
    for c in contacts:
        info = f"• {c.get('name', '?')}"
        if c.get("telegram"):
            info += f" — @{c['telegram']}"
        if c.get("phone"):
            info += f" | {c['phone']}"
        if c.get("email"):
            info += f" | {c['email']}"
        lines.append(info)

    return ToolResult("list_contacts", True, "\n".join(lines))


async def tool_send_telegram(
    username: str = "",
    chat_id: int = 0,
    contact_name: str = "",
    message: str = "",
    db_session=None,
    **kwargs,
) -> ToolResult:
    """
    Отправить сообщение в Telegram.
    v4: Telethon primary для username, Bot API для chat_id.
    """
    if not message:
        return ToolResult("send_telegram", False, "", error="Нужен текст сообщения")

    # Определяем получателя
    recipient_id = None
    recipient_username = ""
    recipient_label = ""

    if chat_id:
        recipient_id = int(chat_id)
        recipient_label = str(chat_id)
    elif contact_name:
        # Smart resolve через ContactBook
        from pds_ultimate.core.contact_book import contact_book
        resolved = contact_book.resolve(contact_name, db_session=db_session)
        if resolved:
            if resolved.get("telegram_id"):
                recipient_id = resolved["telegram_id"]
                recipient_label = f"{resolved.get('name', contact_name)} ({recipient_id})"
            elif resolved.get("telegram"):
                recipient_username = resolved["telegram"].lstrip("@")
                recipient_label = f"@{recipient_username}"
            else:
                return ToolResult(
                    "send_telegram", False, "",
                    error=f"Контакт '{contact_name}' найден, но нет Telegram данных. "
                    f"Привяжи: 'запомни что у {contact_name} телеграм @username'",
                )
        else:
            return ToolResult(
                "send_telegram", False, "",
                error=f"Контакт '{contact_name}' не найден. Укажи username напрямую.",
            )
    elif username:
        recipient_username = username.lstrip("@")
        recipient_label = f"@{recipient_username}"

    # Strategy 1: Telethon userbot — работает по username и chat_id
    try:
        from pds_ultimate.integrations.telethon_client import telethon_client

        if telethon_client._started and telethon_client._client:
            target = recipient_id or recipient_username
            if target:
                await telethon_client._client.send_message(target, message)
                return ToolResult(
                    "send_telegram", True,
                    f"✅ Сообщение отправлено ({recipient_label or target}):\n«{message[:200]}»",
                )
    except Exception as e:
        logger.warning(f"Telethon send failed: {e}")

    # Strategy 2: Bot API fallback (только с chat_id)
    if recipient_id:
        try:
            from pds_ultimate.bot.setup import bot as tg_bot
            if tg_bot:
                await tg_bot.send_message(chat_id=recipient_id, text=message)
                return ToolResult(
                    "send_telegram", True,
                    f"✅ Сообщение отправлено через Bot API ({recipient_label}):\n«{message[:200]}»",
                )
        except Exception as e:
            logger.warning(f"Bot API send failed: {e}")

    # Both failed
    if not recipient_id and not recipient_username:
        return ToolResult(
            "send_telegram", False, "",
            error="Нужен username или chat_id получателя.",
        )

    return ToolResult(
        "send_telegram", False, "",
        error=f"Не удалось отправить сообщение ({recipient_label}). "
        f"Telethon userbot может быть не запущен.",
    )


async def tool_send_email(
    to: str = "",
    subject: str = "",
    body: str = "",
    contact_name: str = "",
    db_session=None,
    **kwargs,
) -> ToolResult:
    """Отправить email через Gmail API или SMTP fallback."""
    if not body:
        return ToolResult("send_email", False, "", error="Нужен текст письма (body)")

    # Определяем получателя
    if not to and contact_name:
        from pds_ultimate.core.contact_book import contact_book
        resolved_email = contact_book.resolve_email(
            contact_name, db_session=db_session)
        if resolved_email:
            to = resolved_email
        else:
            return ToolResult(
                "send_email", False, "",
                error=f"Контакт '{contact_name}' не найден или нет email. "
                f"Привяжи: 'запомни что email {contact_name} — user@example.com'",
            )

    if not to:
        return ToolResult("send_email", False, "", error="Укажи email получателя (to)")

    if not subject:
        subject = "Без темы"

    # ─── Способ 1: Gmail API (если подключён) ───────────────────────
    try:
        from pds_ultimate.integrations.gmail import gmail_client
        if gmail_client._started:
            result = await gmail_client.send_email(to=to, subject=subject, body=body)
            if not result.get("error"):
                return ToolResult(
                    "send_email", True,
                    f"✅ Письмо отправлено (Gmail API) → {to}\nТема: {subject}",
                )
            logger.warning(
                f"Gmail API ошибка: {result['error']}, пробуем SMTP...")
    except Exception as e:
        logger.warning(f"Gmail API недоступен: {e}, пробуем SMTP...")

    # ─── Способ 2: SMTP Fallback ────────────────────────────────────
    try:
        from pds_ultimate.config import config as cfg
        smtp_cfg = cfg.smtp
        if not smtp_cfg.enabled or not smtp_cfg.user or not smtp_cfg.password:
            return ToolResult(
                "send_email", False, "",
                error=(
                    "⚠️ Не удалось отправить email: Gmail API не подключён, "
                    "SMTP не настроен.\n"
                    "Для настройки SMTP:\n"
                    "1. Включите 2FA в Google аккаунте\n"
                    "2. Создайте App Password: https://myaccount.google.com/apppasswords\n"
                    "3. В .env: SMTP_ENABLED=true, SMTP_USER=ваш@gmail.com, SMTP_PASSWORD=ваш_app_password"
                ),
            )

        import asyncio
        import smtplib
        from email.mime.multipart import MIMEMultipart
        from email.mime.text import MIMEText

        def _send_smtp():
            msg = MIMEMultipart("alternative")
            msg["From"] = f"{smtp_cfg.from_name} <{smtp_cfg.user}>"
            msg["To"] = to
            msg["Subject"] = subject
            msg.attach(MIMEText(body, "plain", "utf-8"))
            # HTML версия
            html_body = body.replace("\n", "<br>")
            msg.attach(
                MIMEText(f"<html><body>{html_body}</body></html>", "html", "utf-8"))

            if smtp_cfg.use_tls:
                server = smtplib.SMTP(smtp_cfg.host, smtp_cfg.port, timeout=15)
                server.ehlo()
                server.starttls()
            else:
                server = smtplib.SMTP_SSL(
                    smtp_cfg.host, smtp_cfg.port, timeout=15)

            server.login(smtp_cfg.user, smtp_cfg.password)
            server.send_message(msg)
            server.quit()

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _send_smtp)

        return ToolResult(
            "send_email", True,
            f"✅ Письмо отправлено (SMTP) → {to}\nТема: {subject}",
        )

    except Exception as e:
        err_str = str(e)
        if "SMTPAuthenticationError" in type(e).__name__ or "535" in err_str:
            return ToolResult(
                "send_email", False, "",
                error="❌ SMTP: неверный пароль. Используйте App Password из https://myaccount.google.com/apppasswords",
            )
        return ToolResult("send_email", False, "", error=f"❌ Ошибка отправки email: {e}")


async def tool_get_emails(
    account: str = "",
    max_results: int = 5,
    **kwargs,
) -> ToolResult:
    """Получить непрочитанные письма из Gmail."""
    from pds_ultimate.integrations.gmail import gmail_client

    if not gmail_client._started:
        try:
            await gmail_client.start()
        except Exception as e:
            return ToolResult("get_emails", False, "", error=f"Gmail не подключён: {e}")

    emails = await gmail_client.get_unread(max_results=max_results, account=account or None)
    if not emails:
        return ToolResult("get_emails", True, "📭 Нет непрочитанных писем.")

    lines = [f"📬 Непрочитанных: {len(emails)}\n"]
    for i, em in enumerate(emails, 1):
        lines.append(
            f"{i}. От: {em.get('from', '?')}\n"
            f"   Тема: {em.get('subject', '?')}\n"
            f"   Дата: {em.get('date', '?')}\n"
            f"   Превью: {em.get('snippet', '')[:100]}...\n"
        )
    return ToolResult("get_emails", True, "\n".join(lines), data={"emails": emails})


async def tool_create_file(
    description: str = "",
    file_format: str = "",
    **kwargs,
) -> ToolResult:
    """Создать файл (Excel, Word, PDF, CSV, TXT) по описанию и отправить пользователю."""
    if not description:
        return ToolResult("create_file", False, "", error="Опиши что создать (description)")

    try:
        import json as json_mod
        import os
        from datetime import datetime

        from pds_ultimate.config import USER_FILES_DIR
        from pds_ultimate.core.llm_engine import llm_engine

        os.makedirs(str(USER_FILES_DIR), exist_ok=True)

        # Определяем формат если не указан
        if not file_format:
            fmt_lower = description.lower()
            if any(w in fmt_lower for w in ["excel", "таблиц", "xlsx", "эксель"]):
                file_format = "xlsx"
            elif any(w in fmt_lower for w in ["word", "документ", "docx"]):
                file_format = "docx"
            elif any(w in fmt_lower for w in ["pdf"]):
                file_format = "pdf"
            elif any(w in fmt_lower for w in ["csv"]):
                file_format = "csv"
            else:
                file_format = "xlsx"

        # Генерируем структуру через LLM
        prompt = (
            f"Создай структуру для файла формата {file_format} по запросу: «{description}».\n"
            f"Верни JSON: {{\"title\": \"...\", \"headers\": [\"col1\", ...], \"rows\": [[\"val1\", ...], ...]}}\n"
            f"Добавь примерные данные (5-10 строк). Только JSON, без объяснений."
        )

        raw = await llm_engine.chat(message=prompt, task_type="general", temperature=0.5, json_mode=True)
        try:
            structure = json_mod.loads(raw)
        except Exception:
            structure = {"title": description[:50], "headers": [
                "Данные"], "rows": [["Пример"]]}

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = structure.get("title", "doc").replace(" ", "_")[:30]
        filename = f"{ts}_{safe_title}.{file_format}"
        filepath = str(USER_FILES_DIR / filename)

        if file_format == "xlsx":
            from pds_ultimate.modules.files.excel_engine import ExcelEngine
            engine = ExcelEngine()
            await engine.create(filepath, structure)
        elif file_format == "csv":
            import csv
            with open(filepath, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                writer.writerow(structure.get("headers", []))
                for row in structure.get("rows", []):
                    writer.writerow(row)
        elif file_format == "docx":
            from docx import Document
            doc = Document()
            doc.add_heading(structure.get("title", description[:50]), 0)
            if structure.get("headers"):
                table = doc.add_table(rows=1, cols=len(structure["headers"]))
                table.style = "Table Grid"
                for i, h in enumerate(structure["headers"]):
                    table.rows[0].cells[i].text = str(h)
                for row_data in structure.get("rows", []):
                    row = table.add_row()
                    for i, cell in enumerate(row_data):
                        if i < len(row.cells):
                            row.cells[i].text = str(cell)
                doc.save(filepath)
        elif file_format == "pdf":
            from fpdf import FPDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Helvetica", size=14)
            pdf.cell(200, 10, txt=structure.get(
                "title", "Document"), ln=True, align="C")
            pdf.set_font("Helvetica", size=10)
            for row in structure.get("rows", []):
                pdf.cell(200, 8, txt=" | ".join(str(c) for c in row), ln=True)
            pdf.output(filepath)
        else:
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(structure.get("title", "") + "\n\n")
                for row in structure.get("rows", []):
                    f.write("\t".join(str(c) for c in row) + "\n")

        return ToolResult(
            "create_file", True,
            f"✅ Файл создан: {filename}",
            data={"filepath": filepath, "filename": filename, "send_file": True},
        )
    except Exception as e:
        return ToolResult("create_file", False, "", error=f"Ошибка создания файла: {e}")


async def tool_create_excel(
    title: str = "Таблица",
    headers: str = "",
    rows: str = "",
    **kwargs,
) -> ToolResult:
    """Создать Excel файл с данными и отправить пользователю."""
    try:
        import os
        from datetime import datetime

        from pds_ultimate.config import USER_FILES_DIR
        from pds_ultimate.modules.files.excel_engine import ExcelEngine

        engine = ExcelEngine()

        # Parse headers and rows
        header_list = [h.strip() for h in headers.split(",") if h.strip()] if headers else [
            "Колонка 1", "Колонка 2", "Колонка 3"]
        row_list = []
        if rows:
            for row_str in rows.split(";"):
                cells = [c.strip() for c in row_str.split(",")]
                row_list.append(cells)
        else:
            # Примерные данные
            row_list = [
                ["Пример 1", "Значение A", "100"],
                ["Пример 2", "Значение B", "200"],
                ["Пример 3", "Значение C", "300"],
            ]

        structure = {
            "title": title,
            "headers": header_list,
            "rows": row_list,
        }

        os.makedirs(str(USER_FILES_DIR), exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_title = title.replace(" ", "_")[:30]
        filename = f"{ts}_{safe_title}.xlsx"
        filepath = str(USER_FILES_DIR / filename)

        result = await engine.create(filepath, structure)
        if result.get("success") or result.get("filepath"):
            return ToolResult(
                "create_excel", True,
                f"✅ Excel файл создан: {filename}",
                data={"filepath": filepath,
                      "filename": filename, "send_file": True},
            )
        return ToolResult("create_excel", False, "", error=f"Ошибка: {result}")
    except Exception as e:
        return ToolResult("create_excel", False, "", error=f"Ошибка создания Excel: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# TELEGRAM BOT API — CHAT MANAGEMENT (без Telethon, без my.telegram.org)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_telegram_get_chat_info(
    chat_id: int = 0,
    **kwargs,
) -> ToolResult:
    """Получить информацию о чате через Bot API."""
    if not chat_id:
        return ToolResult("telegram_get_chat_info", False, "", error="Укажи chat_id")

    try:
        from pds_ultimate.bot.setup import bot as tg_bot
        if not tg_bot:
            return ToolResult("telegram_get_chat_info", False, "", error="Бот не инициализирован")

        chat = await tg_bot.get_chat(chat_id=int(chat_id))
        lines = [
            "💬 Информация о чате:",
            f"  🆔 ID: {chat.id}",
            f"  📋 Тип: {chat.type}",
        ]
        if chat.title:
            lines.append(f"  📌 Название: {chat.title}")
        if chat.username:
            lines.append(f"  👤 Username: @{chat.username}")
        if chat.first_name:
            lines.append(f"  🧑 Имя: {chat.first_name} {chat.last_name or ''}")
        if chat.bio:
            lines.append(f"  📝 Био: {chat.bio}")
        if chat.description:
            lines.append(f"  📄 Описание: {chat.description[:200]}")

        return ToolResult(
            "telegram_get_chat_info", True, "\n".join(lines),
            data={
                "id": chat.id,
                "type": chat.type,
                "title": chat.title,
                "username": chat.username,
                "first_name": chat.first_name,
                "last_name": chat.last_name,
            },
        )
    except Exception as e:
        return ToolResult("telegram_get_chat_info", False, "", error=f"Ошибка: {e}")


async def tool_telegram_forward_message(
    from_chat_id: int = 0,
    to_chat_id: int = 0,
    message_id: int = 0,
    **kwargs,
) -> ToolResult:
    """Переслать сообщение из одного чата в другой через Bot API."""
    if not all([from_chat_id, to_chat_id, message_id]):
        return ToolResult(
            "telegram_forward_message", False, "",
            error="Нужны from_chat_id, to_chat_id и message_id",
        )

    try:
        from pds_ultimate.bot.setup import bot as tg_bot
        if not tg_bot:
            return ToolResult("telegram_forward_message", False, "", error="Бот не инициализирован")

        result = await tg_bot.forward_message(
            chat_id=int(to_chat_id),
            from_chat_id=int(from_chat_id),
            message_id=int(message_id),
        )
        return ToolResult(
            "telegram_forward_message", True,
            f"✅ Сообщение переслано: {from_chat_id} → {to_chat_id} (msg_id: {result.message_id})",
        )
    except Exception as e:
        return ToolResult("telegram_forward_message", False, "", error=f"Ошибка: {e}")


async def tool_telegram_pin_message(
    chat_id: int = 0,
    message_id: int = 0,
    **kwargs,
) -> ToolResult:
    """Закрепить сообщение в чате через Bot API."""
    if not chat_id or not message_id:
        return ToolResult("telegram_pin_message", False, "", error="Нужны chat_id и message_id")

    try:
        from pds_ultimate.bot.setup import bot as tg_bot
        if not tg_bot:
            return ToolResult("telegram_pin_message", False, "", error="Бот не инициализирован")

        await tg_bot.pin_chat_message(
            chat_id=int(chat_id),
            message_id=int(message_id),
            disable_notification=True,
        )
        return ToolResult(
            "telegram_pin_message", True,
            f"📌 Сообщение закреплено в чате {chat_id}",
        )
    except Exception as e:
        return ToolResult("telegram_pin_message", False, "", error=f"Ошибка: {e}")


async def tool_telegram_manage_chat(
    action: str = "info",
    chat_id: int = 0,
    user_id: int = 0,
    title: str = "",
    description: str = "",
    **kwargs,
) -> ToolResult:
    """
    Управление чатом/группой через Bot API.
    Действия: info, set_title, set_description, ban, unban, get_members_count.
    Бот должен быть админом группы.
    """
    if not chat_id:
        return ToolResult("telegram_manage_chat", False, "", error="Укажи chat_id")

    try:
        from pds_ultimate.bot.setup import bot as tg_bot
        if not tg_bot:
            return ToolResult("telegram_manage_chat", False, "", error="Бот не инициализирован")

        if action == "info":
            chat = await tg_bot.get_chat(chat_id=int(chat_id))
            count = await tg_bot.get_chat_member_count(chat_id=int(chat_id))
            lines = [
                f"💬 Чат: {chat.title or chat.first_name or chat_id}",
                f"  🆔 ID: {chat.id}",
                f"  📋 Тип: {chat.type}",
                f"  👥 Участников: {count}",
            ]
            if chat.description:
                lines.append(f"  📄 Описание: {chat.description[:200]}")
            return ToolResult("telegram_manage_chat", True, "\n".join(lines))

        elif action == "set_title" and title:
            await tg_bot.set_chat_title(chat_id=int(chat_id), title=title)
            return ToolResult("telegram_manage_chat", True, f"✅ Название изменено: {title}")

        elif action == "set_description":
            await tg_bot.set_chat_description(chat_id=int(chat_id), description=description)
            return ToolResult("telegram_manage_chat", True, "✅ Описание обновлено")

        elif action == "ban" and user_id:
            await tg_bot.ban_chat_member(chat_id=int(chat_id), user_id=int(user_id))
            return ToolResult("telegram_manage_chat", True, f"🚫 Пользователь {user_id} заблокирован")

        elif action == "unban" and user_id:
            await tg_bot.unban_chat_member(chat_id=int(chat_id), user_id=int(user_id), only_if_banned=True)
            return ToolResult("telegram_manage_chat", True, f"✅ Пользователь {user_id} разблокирован")

        elif action == "get_members_count":
            count = await tg_bot.get_chat_member_count(chat_id=int(chat_id))
            return ToolResult("telegram_manage_chat", True, f"👥 Участников: {count}")

        else:
            return ToolResult(
                "telegram_manage_chat", False, "",
                error=f"Неизвестное действие: {action}. Доступны: info, set_title, set_description, ban, unban, get_members_count",
            )

    except Exception as e:
        return ToolResult("telegram_manage_chat", False, "", error=f"Ошибка: {e}")


async def tool_telegram_send_photo(
    chat_id: int = 0,
    photo_path: str = "",
    caption: str = "",
    **kwargs,
) -> ToolResult:
    """Отправить фото в Telegram через Bot API."""
    if not chat_id or not photo_path:
        return ToolResult("telegram_send_photo", False, "", error="Нужны chat_id и photo_path")

    try:
        import os

        from aiogram.types import FSInputFile

        from pds_ultimate.bot.setup import bot as tg_bot
        if not tg_bot:
            return ToolResult("telegram_send_photo", False, "", error="Бот не инициализирован")

        if not os.path.exists(photo_path):
            return ToolResult("telegram_send_photo", False, "", error=f"Файл не найден: {photo_path}")

        photo = FSInputFile(photo_path)
        await tg_bot.send_photo(chat_id=int(chat_id), photo=photo, caption=caption or None)
        return ToolResult(
            "telegram_send_photo", True,
            f"📷 Фото отправлено в чат {chat_id}",
        )
    except Exception as e:
        return ToolResult("telegram_send_photo", False, "", error=f"Ошибка: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# SANDBOX TOOLS (handlers) — Safe file operations
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_sandbox_read_file(path: str, start_line: int = 0, end_line: int = 0, **kwargs) -> ToolResult:
    """Read file via sandbox engine."""
    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        result = sandbox.read_file(
            path,
            start_line=int(start_line) if start_line else None,
            end_line=int(end_line) if end_line else None,
        )
        return ToolResult("sandbox_read_file", True, result)
    except Exception as e:
        return ToolResult("sandbox_read_file", False, "", error=str(e))


async def tool_sandbox_edit_file(
    path: str, edits: str, create_backup: bool = True, **kwargs
) -> ToolResult:
    """Edit file safely with backup + AST validation."""
    import json as _json

    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        edits_list = _json.loads(edits) if isinstance(edits, str) else edits
        result = sandbox.edit_file(
            path, edits_list, create_backup=bool(create_backup))
        return ToolResult("sandbox_edit_file", True, result)
    except _json.JSONDecodeError:
        return ToolResult("sandbox_edit_file", False, "",
                          error="edits должен быть валидный JSON массив")
    except Exception as e:
        return ToolResult("sandbox_edit_file", False, "", error=str(e))


async def tool_sandbox_create_file(path: str, content: str, **kwargs) -> ToolResult:
    """Create file with syntax validation."""
    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        result = sandbox.create_file(path, content)
        return ToolResult("sandbox_create_file", True, result)
    except Exception as e:
        return ToolResult("sandbox_create_file", False, "", error=str(e))


async def tool_sandbox_run_code(code: str, timeout: int = 30, **kwargs) -> ToolResult:
    """Execute Python code in sandbox."""
    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        result = sandbox.execute_code(code, timeout=int(timeout))
        return ToolResult("sandbox_run_code", True, result)
    except Exception as e:
        return ToolResult("sandbox_run_code", False, "", error=str(e))


async def tool_sandbox_search_files(
    pattern: str, directory: str = "", extensions: str = "",
    regex: bool = False, **kwargs
) -> ToolResult:
    """Search files by pattern (grep-like)."""
    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        ext_list = [e.strip() for e in extensions.split(
            ",") if e.strip()] if extensions else None
        result = sandbox.search_in_files(
            pattern, directory=directory or None,
            extensions=ext_list, regex=bool(regex),
        )
        return ToolResult("sandbox_search_files", True, result)
    except Exception as e:
        return ToolResult("sandbox_search_files", False, "", error=str(e))


async def tool_sandbox_list_dir(path: str = "", max_depth: int = 3, **kwargs) -> ToolResult:
    """List directory tree."""
    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        result = sandbox.list_directory(path or None, max_depth=int(max_depth))
        return ToolResult("sandbox_list_dir", True, result)
    except Exception as e:
        return ToolResult("sandbox_list_dir", False, "", error=str(e))


async def tool_sandbox_csv_read(path: str, max_rows: int = 30, **kwargs) -> ToolResult:
    """Read CSV with formatted table."""
    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        result = sandbox.read_csv(path, max_rows=int(max_rows))
        return ToolResult("sandbox_csv_read", True, result)
    except Exception as e:
        return ToolResult("sandbox_csv_read", False, "", error=str(e))


async def tool_sandbox_csv_edit(path: str, operations: str, **kwargs) -> ToolResult:
    """Edit CSV file."""
    import json as _json

    from pds_ultimate.core.sandbox_engine import sandbox
    try:
        ops = _json.loads(operations) if isinstance(
            operations, str) else operations
        result = sandbox.edit_csv(path, ops)
        return ToolResult("sandbox_csv_edit", True, result)
    except _json.JSONDecodeError:
        return ToolResult("sandbox_csv_edit", False, "",
                          error="operations должен быть валидный JSON массив")
    except Exception as e:
        return ToolResult("sandbox_csv_edit", False, "", error=str(e))


# ═══════════════════════════════════════════════════════════════════════════════
# WIDE RESEARCH TOOLS (handlers) — Parallel sub-agents
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_wide_research(query: str, max_sources: int = 5, **kwargs) -> ToolResult:
    """Wide research with parallel sub-agents."""
    from pds_ultimate.core.wide_research import wide_research
    try:
        report = await wide_research.research(
            query=query,
            max_sources_per_query=int(max_sources),
        )
        return ToolResult(
            "wide_research", True, report.summary(),
            data={
                "total_findings": len(report.findings),
                "contradictions": len(report.contradictions),
                "insights": report.insights[:3] if report.insights else [],
                "confidence": report.overall_confidence,
            },
        )
    except Exception as e:
        return ToolResult("wide_research", False, "", error=f"Ошибка исследования: {e}")


async def tool_quick_research(query: str, max_sources: int = 3, **kwargs) -> ToolResult:
    """Quick research without LLM."""
    from pds_ultimate.core.wide_research import wide_research
    try:
        report = await wide_research.quick_research(
            query=query, max_sources=int(max_sources),
        )
        return ToolResult(
            "quick_research_v2", True, report.summary(),
            data={
                "findings": len(report.findings),
                "confidence": report.overall_confidence,
            },
        )
    except Exception as e:
        return ToolResult("quick_research_v2", False, "", error=str(e))


async def tool_compare_research(items: str, criteria: str = "", **kwargs) -> ToolResult:
    """Compare N items on M criteria via parallel research."""
    from pds_ultimate.core.wide_research import wide_research
    try:
        items_list = [i.strip() for i in items.split(",") if i.strip()]
        criteria_list = [c.strip() for c in criteria.split(
            ",") if c.strip()] if criteria else None
        if len(items_list) < 2:
            return ToolResult("compare_research", False, "",
                              error="Нужно минимум 2 объекта для сравнения (через запятую)")
        report = await wide_research.compare_research(
            items=items_list, criteria=criteria_list,
        )
        return ToolResult(
            "compare_research", True, report.summary(),
            data={"items": items_list, "findings": len(report.findings)},
        )
    except Exception as e:
        return ToolResult("compare_research", False, "", error=str(e))


# ═══════════════════════════════════════════════════════════════════════════════
# DATA ANALYSIS TOOLS (handlers) — Built-in analytics
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_analyze_data(path: str, generate_charts: bool = True, **kwargs) -> ToolResult:
    """Full EDA on a data file."""
    from pds_ultimate.core.data_analysis import data_engine
    try:
        result = await data_engine.eda(path, generate_charts=bool(generate_charts))
        if not result.success:
            return ToolResult("analyze_data", False, "", error=result.error)
        data = result.data.copy()
        if result.charts:
            data["charts"] = result.charts
            data["send_file"] = True
            data["filepath"] = result.charts[0]
            data["filename"] = os.path.basename(result.charts[0])
        return ToolResult("analyze_data", True, result.full_summary(), data=data)
    except Exception as e:
        return ToolResult("analyze_data", False, "", error=str(e))


async def tool_create_chart(
    path: str, x_column: str, y_column: str,
    chart_type: str = "bar", title: str = "", **kwargs
) -> ToolResult:
    """Create chart from data file."""
    from pds_ultimate.core.data_analysis import data_engine
    try:
        result = await data_engine.generate_chart(
            path, x_column, y_column,
            chart_type=chart_type, title=title,
        )
        if not result.success:
            return ToolResult("create_chart", False, "", error=result.error)
        data = result.data.copy()
        if result.charts:
            data["send_file"] = True
            data["filepath"] = result.charts[0]
            data["filename"] = os.path.basename(result.charts[0])
        return ToolResult("create_chart", True, result.full_summary(), data=data)
    except Exception as e:
        return ToolResult("create_chart", False, "", error=str(e))


async def tool_data_filter(
    path: str, column: str, condition: str, value: str = "", **kwargs
) -> ToolResult:
    """Filter data by condition."""
    from pds_ultimate.core.data_analysis import data_engine
    try:
        result = await data_engine.filter_data(path, column, condition, value)
        if not result.success:
            return ToolResult("data_filter", False, "", error=result.error)
        return ToolResult("data_filter", True, result.full_summary(), data=result.data)
    except Exception as e:
        return ToolResult("data_filter", False, "", error=str(e))


async def tool_data_group_by(
    path: str, group_column: str,
    agg_column: str = "", agg_func: str = "count", **kwargs
) -> ToolResult:
    """Group by with aggregation."""
    from pds_ultimate.core.data_analysis import data_engine
    try:
        result = await data_engine.group_by(path, group_column, agg_column, agg_func)
        if not result.success:
            return ToolResult("data_group_by", False, "", error=result.error)
        data = result.data.copy()
        if result.charts:
            data["send_file"] = True
            data["filepath"] = result.charts[0]
            data["filename"] = os.path.basename(result.charts[0])
        return ToolResult("data_group_by", True, result.full_summary(), data=data)
    except Exception as e:
        return ToolResult("data_group_by", False, "", error=str(e))


async def tool_data_stats(path: str, column: str = "", **kwargs) -> ToolResult:
    """Detailed statistics."""
    from pds_ultimate.core.data_analysis import data_engine
    try:
        result = await data_engine.compute_stats(path, column=column)
        if not result.success:
            return ToolResult("data_stats", False, "", error=result.error)
        return ToolResult("data_stats", True, result.full_summary(), data=result.data)
    except Exception as e:
        return ToolResult("data_stats", False, "", error=str(e))


# ═══════════════════════════════════════════════════════════════════════════════
# РЕГИСТРАЦИЯ ВСЕХ TOOLS
# ═══════════════════════════════════════════════════════════════════════════════

def register_all_tools() -> int:
    """
    Зарегистрировать все бизнес-инструменты.
    Вызывается при старте системы.

    Returns:
        Количество зарегистрированных tools.
    """
    tools = [
        # ─── Логистика ───────────────────────────────────────────────
        Tool(
            name="create_order",
            description="Создать новый заказ. Принимает текстовое описание позиций товаров.",
            parameters=[
                ToolParameter("items_text", "string",
                              "Текст с позициями (название, количество, единица, цена)", True),
            ],
            handler=tool_create_order,
            category="logistics",
            needs_db=True,
        ),
        Tool(
            name="get_orders_status",
            description="Получить статус заказа или список всех активных заказов.",
            parameters=[
                ToolParameter("order_number", "string",
                              "Номер заказа (например ORD-0001). Если не указан — все активные.", False),
            ],
            handler=tool_get_orders_status,
            category="logistics",
            needs_db=True,
        ),
        Tool(
            name="set_income",
            description="Установить доход (сколько заплатили МНЕ) за заказ.",
            parameters=[
                ToolParameter("order_number", "string", "Номер заказа", True),
                ToolParameter("amount", "number", "Сумма дохода", True),
                ToolParameter("currency", "string",
                              "Валюта (USD/CNY/TMT)", False, "USD"),
            ],
            handler=tool_set_income,
            category="finance",
            needs_db=True,
        ),
        Tool(
            name="set_expense",
            description="Установить расход на товар (сколько Я заплатил поставщику).",
            parameters=[
                ToolParameter("order_number", "string", "Номер заказа", True),
                ToolParameter("amount", "number", "Сумма расхода", True),
                ToolParameter("currency", "string",
                              "Валюта (USD/CNY/TMT)", False, "USD"),
            ],
            handler=tool_set_expense,
            category="finance",
            needs_db=True,
        ),

        # ─── Финансы ─────────────────────────────────────────────────
        Tool(
            name="get_financial_summary",
            description="Получить полную финансовую сводку: доходы, расходы, прибыль, баланс.",
            parameters=[],
            handler=tool_get_financial_summary,
            category="finance",
            needs_db=True,
        ),
        Tool(
            name="convert_currency",
            description="Конвертировать валюту. Фиксированные курсы: 1 USD = 19.5 TMT, 1 USD = 7.1 CNY.",
            parameters=[
                ToolParameter("amount", "number", "Сумма", True),
                ToolParameter("from_currency", "string",
                              "Из какой валюты (USD/CNY/TMT)", True),
                ToolParameter("to_currency", "string",
                              "В какую валюту", False, "USD"),
            ],
            handler=tool_convert_currency,
            category="finance",
        ),

        # ─── Контакты ───────────────────────────────────────────────
        Tool(
            name="save_contact_note",
            description="Сохранить заметку или предупреждение о контрагенте/контакте.",
            parameters=[
                ToolParameter("name", "string", "Имя контакта", True),
                ToolParameter("note", "string", "Текст заметки", True),
                ToolParameter("is_warning", "boolean",
                              "Это предупреждение?", False, False),
            ],
            handler=tool_save_contact_note,
            category="contacts",
            needs_db=True,
        ),
        Tool(
            name="find_contact",
            description="Найти контакт по имени. Показывает заметки и предупреждения.",
            parameters=[
                ToolParameter("query", "string",
                              "Имя или часть имени контакта", True),
            ],
            handler=tool_find_contact,
            category="contacts",
            needs_db=True,
        ),

        # ─── Календарь ──────────────────────────────────────────────
        Tool(
            name="create_reminder",
            description="Создать напоминание на определённую дату и время.",
            parameters=[
                ToolParameter("message", "string", "Текст напоминания", True),
                ToolParameter("scheduled_at", "string",
                              "Дата и время (формат: YYYY-MM-DD HH:MM или DD.MM.YYYY HH:MM)", True),
            ],
            handler=tool_create_reminder,
            category="calendar",
            needs_db=True,
        ),
        Tool(
            name="create_calendar_event",
            description="Создать событие в календаре.",
            parameters=[
                ToolParameter("title", "string", "Название события", True),
                ToolParameter("event_date", "string",
                              "Дата и время (формат: YYYY-MM-DD HH:MM)", True),
                ToolParameter("description", "string",
                              "Описание события", False, ""),
            ],
            handler=tool_create_calendar_event,
            category="calendar",
            needs_db=True,
        ),

        # ─── Отчёты ─────────────────────────────────────────────────
        Tool(
            name="morning_brief",
            description="Сформировать утренний брифинг с обзором заказов, позиций и финансов.",
            parameters=[],
            handler=tool_morning_brief,
            category="reports",
            needs_db=True,
        ),

        # ─── Текст ──────────────────────────────────────────────────
        Tool(
            name="translate",
            description="Перевести текст на другой язык.",
            parameters=[
                ToolParameter("text", "string", "Текст для перевода", True),
                ToolParameter("target_lang", "string",
                              "Целевой язык (ru/en/zh/tr)", False, "ru"),
                ToolParameter("source_lang", "string", "Исходный язык", False),
            ],
            handler=tool_translate,
            category="text",
        ),
        Tool(
            name="summarize",
            description="Создать краткое саммари текста.",
            parameters=[
                ToolParameter("text", "string",
                              "Текст для суммаризации", True),
            ],
            handler=tool_summarize,
            category="text",
        ),

        # ─── Безопасность ────────────────────────────────────────────
        Tool(
            name="security_emergency",
            description="ЭКСТРЕННОЕ УДАЛЕНИЕ финансовых данных. Только по кодовому слову!",
            parameters=[],
            handler=tool_security_emergency,
            category="security",
            needs_db=True,
            visible=False,  # Не показывать в system prompt
        ),

        # ─── Память ─────────────────────────────────────────────────
        Tool(
            name="remember",
            description="Запомнить важный факт, предпочтение или правило для будущего использования.",
            parameters=[
                ToolParameter("fact", "string", "Что запомнить", True),
                ToolParameter("importance", "number",
                              "Важность от 0.0 до 1.0", False, 0.5),
                ToolParameter("memory_type", "string",
                              "Тип: fact/preference/rule/knowledge", False, "fact"),
            ],
            handler=tool_remember,
            category="memory",
        ),
        Tool(
            name="recall",
            description="Вспомнить факты из долгосрочной памяти по ключевым словам.",
            parameters=[
                ToolParameter("query", "string", "Что вспомнить", True),
            ],
            handler=tool_recall,
            category="memory",
        ),

        # ─── Браузер (Manus-level) ────────────────────────────────────
        Tool(
            name="web_search",
            description=(
                "Поиск в интернете через DuckDuckGo. Возвращает список "
                "результатов (заголовок, URL, сниппет). Используй для поиска "
                "информации, цен, поставщиков, новостей, курсов."
            ),
            parameters=[
                ToolParameter("query", "string", "Поисковый запрос", True),
                ToolParameter("max_results", "number",
                              "Максимум результатов (1-20)", False, 10),
            ],
            handler=tool_web_search,
            category="browser",
        ),
        Tool(
            name="open_page",
            description=(
                "Открыть веб-страницу и извлечь содержимое "
                "(текст, ссылки, таблицы, мета-данные). "
                "Используй после web_search чтобы прочитать конкретную страницу."
            ),
            parameters=[
                ToolParameter("url", "string", "URL страницы", True),
            ],
            handler=tool_open_page,
            category="browser",
        ),
        Tool(
            name="search_and_read",
            description=(
                "Manus-level: Поиск → автоматически открывает топ-N страниц → "
                "извлекает текст со всех. Идеально для быстрого исследования "
                "вопроса. Возвращает контент с нескольких источников сразу."
            ),
            parameters=[
                ToolParameter("query", "string", "Поисковый запрос", True),
                ToolParameter("max_pages", "number",
                              "Сколько страниц открыть (1-5)", False, 3),
            ],
            handler=tool_search_and_read,
            category="browser",
        ),
        Tool(
            name="deep_web_research",
            description=(
                "Глубокое исследование: поиск → открытие страниц → "
                "переход по релевантным ссылкам → сбор данных из всех "
                "источников. Для сложных вопросов где нужно много фактов."
            ),
            parameters=[
                ToolParameter("query", "string", "Тема исследования", True),
                ToolParameter("max_sources", "number",
                              "Макс. источников (1-10)", False, 5),
            ],
            handler=tool_deep_web_research,
            category="browser",
        ),
        Tool(
            name="extract_page_data",
            description=(
                "Извлечь структурированные данные со страницы: "
                "заголовки, таблицы, ссылки, мета-теги. "
                "Можно указать фокус для фильтрации по теме."
            ),
            parameters=[
                ToolParameter("url", "string", "URL страницы", True),
                ToolParameter("focus", "string",
                              "На чём сфокусироваться (опционально)", False),
            ],
            handler=tool_extract_page_data,
            category="browser",
        ),
        Tool(
            name="browser_screenshot",
            description="Сделать скриншот текущей страницы в браузере (Playwright).",
            parameters=[
                ToolParameter("full_page", "boolean",
                              "Полная страница (true) или видимая область", False),
            ],
            handler=tool_browser_screenshot,
            category="browser",
        ),
        Tool(
            name="browser_click",
            description="Кликнуть по элементу на странице (CSS-селектор, Playwright).",
            parameters=[
                ToolParameter("selector", "string",
                              "CSS-селектор элемента", True),
            ],
            handler=tool_browser_click,
            category="browser",
        ),
        Tool(
            name="browser_fill",
            description="Заполнить поле на веб-странице текстом (Playwright).",
            parameters=[
                ToolParameter("selector", "string", "CSS-селектор поля", True),
                ToolParameter("value", "string", "Текст для ввода", True),
            ],
            handler=tool_browser_fill,
            category="browser",
        ),

        # ─── Исследование (Internet Reasoning) ──────────────────────
        Tool(
            name="research",
            description=(
                "Исследовать вопрос с проверкой множества источников. "
                "Ищет в интернете, извлекает факты, оценивает достоверность, "
                "обнаруживает противоречия и синтезирует ответ. "
                "Используй для проверки фактов, сравнения цен, "
                "анализа рынка, поиска информации."
            ),
            parameters=[
                ToolParameter("query", "string",
                              "Вопрос для исследования", True),
                ToolParameter("max_sources", "number",
                              "Максимум источников (1-10)", False, 5),
            ],
            handler=tool_research,
            category="research",
        ),
        Tool(
            name="deep_research",
            description=(
                "Глубокое исследование с максимальным покрытием. "
                "Расширяет запросы, анализирует до 10 источников, "
                "извлекает больше фактов. Для сложных вопросов, "
                "где нужна проверка из множества независимых источников."
            ),
            parameters=[
                ToolParameter("query", "string",
                              "Вопрос для глубокого исследования", True),
                ToolParameter("max_sources", "number",
                              "Максимум источников (1-15)", False, 10),
            ],
            handler=tool_deep_research,
            category="research",
        ),
        Tool(
            name="quick_search",
            description=(
                "Быстрый поиск с анализом — без расширения запросов. "
                "Для простых вопросов, когда нужен быстрый ответ "
                "с оценкой достоверности."
            ),
            parameters=[
                ToolParameter("query", "string",
                              "Поисковый запрос", True),
            ],
            handler=tool_quick_search,
            category="research",
        ),

        # ─── Part 7: Бизнес-интеграции ──────────────────────────────
        Tool(
            name="exchange_rates",
            description=(
                "Получить актуальный курс обмена валют. "
                "Онлайн-курсы + фиксированные (TMT, CNY). "
                "Можно конвертировать сумму между валютами."
            ),
            parameters=[
                ToolParameter("from_currency", "string",
                              "Из какой валюты (USD/CNY/TMT/EUR)", False, "USD"),
                ToolParameter("to_currency", "string",
                              "В какую валюту (если пусто — все курсы)", False),
                ToolParameter("amount", "number",
                              "Сумма для конвертации", False, 1.0),
            ],
            handler=tool_exchange_rates,
            category="finance",
        ),
        Tool(
            name="google_calendar",
            description=(
                "Работа с Google Calendar: просмотр событий на сегодня, "
                "создание новых событий, поиск свободных слотов."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: today/create/free_slots", False, "today"),
                ToolParameter("title", "string",
                              "Название события (для create)", False),
                ToolParameter("start_time", "string",
                              "Начало (YYYY-MM-DD HH:MM)", False),
                ToolParameter("end_time", "string",
                              "Конец (YYYY-MM-DD HH:MM)", False),
                ToolParameter("description", "string",
                              "Описание события", False),
            ],
            handler=tool_google_calendar_events,
            category="calendar",
        ),

        # ─── Part 7: Файловые движки ────────────────────────────────
        Tool(
            name="ocr_recognize",
            description=(
                "Распознать текст на изображении (OCR). "
                "Поддержка: фото чеков, накладных, документов. "
                "Языки: RU, EN, ZH. Может извлечь суммы и трекинг-номера."
            ),
            parameters=[
                ToolParameter("file_path", "string",
                              "Путь к файлу изображения", True),
                ToolParameter("extract_amounts", "boolean",
                              "Извлечь денежные суммы", False, False),
                ToolParameter("extract_tracking", "boolean",
                              "Извлечь трекинг-номера", False, False),
            ],
            handler=tool_ocr_recognize,
            category="files",
        ),
        Tool(
            name="convert_file",
            description=(
                "Конвертировать файл в другой формат. "
                "Поддержка: xlsx↔csv, docx→pdf, pdf→txt, json→csv и другие."
            ),
            parameters=[
                ToolParameter("file_path", "string",
                              "Путь к исходному файлу", True),
                ToolParameter("target_format", "string",
                              "Целевой формат (csv/pdf/xlsx/txt/json)", True),
            ],
            handler=tool_convert_file,
            category="files",
        ),

        # ─── Part 7: Исполнительные инструменты ─────────────────────
        Tool(
            name="scan_receipt",
            description=(
                "Сканировать чек/квитанцию: OCR + распознавание "
                "позиций, итога, категории расхода. "
                "Автоматически сохраняет в базу расходов."
            ),
            parameters=[
                ToolParameter("file_path", "string",
                              "Путь к фото чека", True),
                ToolParameter("save_to_db", "boolean",
                              "Сохранить в базу расходов", False, True),
            ],
            handler=tool_scan_receipt,
            category="finance",
            needs_db=True,
        ),
        Tool(
            name="translate_text",
            description=(
                "Перевести текст с бизнес-глоссарием. "
                "Автоопределение языка. "
                "Поддержка: RU, EN, ZH, TK, TR, AR, FA, DE, FR, ES, IT, PT."
            ),
            parameters=[
                ToolParameter("text", "string", "Текст для перевода", True),
                ToolParameter("target_lang", "string",
                              "Целевой язык (ru/en/zh/tk)", False, "ru"),
                ToolParameter("source_lang", "string",
                              "Исходный язык (авто если пусто)", False),
            ],
            handler=tool_translate_text,
            category="text",
        ),
        Tool(
            name="archivist_rename",
            description=(
                "Стандартизировать имя файла по корпоративному стандарту. "
                "Формат: YYYY_MM_DD_Category_Description.ext. "
                "Автоопределение категории из содержимого."
            ),
            parameters=[
                ToolParameter("file_path", "string",
                              "Путь к файлу", True),
                ToolParameter("description", "string",
                              "Описание файла (опционально)", False),
            ],
            handler=tool_archivist_rename,
            category="files",
        ),

        # ─── Part 8: Plugin System ──────────────────────────────────
        Tool(
            name="plugin_connect",
            description=(
                "Подключить внешний API как плагин. "
                "Автоматически определяет тип API по URL или ключу. "
                "Поддержка: OpenAI, Anthropic, Stripe, SendGrid, Twilio, Google, Telegram и другие."
            ),
            parameters=[
                ToolParameter("name", "string",
                              "Имя плагина (например: 'stripe', 'my_api')", True),
                ToolParameter("base_url", "string",
                              "Базовый URL API", True),
                ToolParameter("api_key", "string",
                              "API ключ (если нужен)", False),
                ToolParameter("plugin_type", "string",
                              "Тип: REST_API/LLM_API/PAYMENT_API/MESSAGING_API/CLOUD_API/WEBHOOK", False, "REST_API"),
            ],
            handler=tool_plugin_connect,
            category="plugins",
        ),
        Tool(
            name="plugin_execute",
            description=(
                "Выполнить действие через подключённый плагин. "
                "Вызывает endpoint API с указанными параметрами."
            ),
            parameters=[
                ToolParameter("plugin_name", "string",
                              "Имя плагина", True),
                ToolParameter("endpoint", "string",
                              "Путь endpoint (например '/chat/completions')", True),
                ToolParameter("method", "string",
                              "HTTP метод (GET/POST/PUT/DELETE)", False, "GET"),
                ToolParameter("body", "string",
                              "Тело запроса (JSON строка)", False),
            ],
            handler=tool_plugin_execute,
            category="plugins",
        ),
        Tool(
            name="plugin_list",
            description="Показать список подключённых плагинов и их статус.",
            parameters=[],
            handler=tool_plugin_list,
            category="plugins",
        ),

        # ─── Part 8: Autonomous Tasks ───────────────────────────────
        Tool(
            name="autonomous_task",
            description=(
                "Создать автономную задачу. Агент декомпозирует цель на шаги "
                "и выполняет их самостоятельно с самокоррекцией при ошибках. "
                "Для сложных многошаговых задач."
            ),
            parameters=[
                ToolParameter("goal", "string",
                              "Описание цели (что нужно сделать)", True),
                ToolParameter("priority", "string",
                              "Приоритет: critical/high/normal/low/background", False, "normal"),
                ToolParameter("deadline_hours", "number",
                              "Дедлайн в часах (0 = без дедлайна)", False, 0),
            ],
            handler=tool_autonomous_task,
            category="autonomy",
        ),
        Tool(
            name="task_status",
            description="Показать статус автономных задач.",
            parameters=[
                ToolParameter("task_id", "string",
                              "ID задачи (если пусто — все активные)", False),
            ],
            handler=tool_task_status,
            category="autonomy",
        ),

        # ─── Part 8: Memory & Learning ──────────────────────────────
        Tool(
            name="learn_skill",
            description=(
                "Научить агента новому навыку/стратегии. "
                "Агент запомнит паттерн и будет использовать его в будущем."
            ),
            parameters=[
                ToolParameter("name", "string", "Название навыка", True),
                ToolParameter("pattern", "string",
                              "Regex паттерн для активации (например 'курс|валют')", True),
                ToolParameter("strategy", "string",
                              "Описание стратегии (что делать)", True),
            ],
            handler=tool_learn_skill,
            category="memory",
        ),
        Tool(
            name="memory_stats",
            description="Статистика памяти: навыки, ошибки, паттерны, обучение.",
            parameters=[],
            handler=tool_memory_stats,
            category="memory",
        ),

        # ─── Part 9: Smart Triggers ─────────────────────────────────
        Tool(
            name="set_trigger",
            description=(
                "Установить умный триггер/алерт. "
                "Типы: exchange_rate (курс), balance (баланс), "
                "supplier_silence (тишина поставщика), deadline, price_change. "
                "Или пользовательский триггер на любое условие."
            ),
            parameters=[
                ToolParameter("name", "string", "Название триггера", True),
                ToolParameter("trigger_type", "string",
                              "Тип: threshold/silence/exchange_rate/balance/deadline/price_change/custom",
                              False, "threshold"),
                ToolParameter("field", "string",
                              "Поле для мониторинга (rate_usd_cny, balance, etc.)", False),
                ToolParameter("operator", "string",
                              "Оператор: >/>=/</<=/==/!=", False, ">"),
                ToolParameter("value", "string",
                              "Пороговое значение", False),
                ToolParameter("severity", "string",
                              "Серьёзность: info/warning/critical/emergency", False, "warning"),
                ToolParameter("template", "string",
                              "Шаблон: exchange_rate/balance/supplier_silence/deadline/price_change",
                              False),
            ],
            handler=tool_set_trigger,
            category="triggers",
        ),
        Tool(
            name="list_triggers",
            description="Показать список активных триггеров и историю алертов.",
            parameters=[
                ToolParameter("show_history", "boolean",
                              "Показать историю алертов", False, False),
            ],
            handler=tool_list_triggers,
            category="triggers",
        ),

        # ─── Part 9: Analytics Dashboard ────────────────────────────
        Tool(
            name="dashboard",
            description=(
                "Бизнес-дашборд: ключевые метрики, KPI, тренды. "
                "Записывает метрики и показывает аналитику."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: show/record/trend/forecast", False, "show"),
                ToolParameter("metric_name", "string",
                              "Имя метрики (для record/trend/forecast)", False),
                ToolParameter("value", "number",
                              "Значение (для record)", False),
                ToolParameter("unit", "string",
                              "Единица измерения", False, ""),
            ],
            handler=tool_dashboard,
            category="analytics",
        ),
        Tool(
            name="kpi_track",
            description=(
                "Отслеживание KPI: создать цель, обновить прогресс, "
                "показать доску KPI."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: create/update/board", False, "board"),
                ToolParameter("name", "string", "Название KPI", False),
                ToolParameter("target", "number", "Целевое значение", False),
                ToolParameter("value", "number",
                              "Текущее значение (для update)", False),
                ToolParameter("unit", "string",
                              "Единица измерения", False, ""),
            ],
            handler=tool_kpi_track,
            category="analytics",
        ),

        # ─── Part 9: CRM ────────────────────────────────────────────
        Tool(
            name="rate_contact",
            description=(
                "Оценить контакт/поставщика (1-5 звёзд). "
                "Можно оценить в целом или по категориям: "
                "reliability, quality, pricing, communication, delivery_speed."
            ),
            parameters=[
                ToolParameter("name", "string",
                              "Имя контакта/поставщика", True),
                ToolParameter("rating", "number",
                              "Рейтинг (1-5 звёзд)", True),
                ToolParameter("comment", "string",
                              "Комментарий к оценке", False, ""),
                ToolParameter("category", "string",
                              "Категория: reliability/quality/pricing/communication/delivery_speed",
                              False, ""),
            ],
            handler=tool_rate_contact,
            category="crm",
        ),
        Tool(
            name="crm_search",
            description=(
                "Поиск в CRM: контакты, сделки, pipeline. "
                "Фильтрация по типу, рейтингу, тегам."
            ),
            parameters=[
                ToolParameter("query", "string",
                              "Поисковый запрос (имя, компания)", False, ""),
                ToolParameter("action", "string",
                              "Действие: search/pipeline/stats/add_contact/add_deal",
                              False, "search"),
                ToolParameter("contact_type", "string",
                              "Тип: supplier/client/partner/logistics/other", False, ""),
                ToolParameter("min_rating", "number",
                              "Минимальный рейтинг (0-5)", False, 0),
            ],
            handler=tool_crm_search,
            category="crm",
        ),

        # ─── Part 9: Evening Digest ─────────────────────────────────
        Tool(
            name="evening_digest",
            description=(
                "Вечерний дайджест: итоги дня, сравнение с вчера, "
                "рекомендации на завтра. Автоматическая аналитика."
            ),
            parameters=[
                ToolParameter("format", "string",
                              "Формат: full/short", False, "full"),
                ToolParameter("revenue", "number",
                              "Доход за сегодня (если не из БД)", False, 0),
                ToolParameter("expenses", "number",
                              "Расходы за сегодня", False, 0),
                ToolParameter("orders_created", "number",
                              "Заказов создано", False, 0),
                ToolParameter("tasks_completed", "number",
                              "Задач завершено", False, 0),
            ],
            handler=tool_evening_digest,
            category="reports",
        ),

        # ─── Part 9: Workflow & Templates ────────────────────────────
        Tool(
            name="create_template",
            description=(
                "Создать шаблон заказа, чек-лист или workflow. "
                "Шаблоны можно переиспользовать для быстрого создания."
            ),
            parameters=[
                ToolParameter("name", "string", "Название шаблона", True),
                ToolParameter("template_type", "string",
                              "Тип: order/checklist/workflow/message", False, "checklist"),
                ToolParameter("content", "string",
                              "Содержимое/шаги (каждый шаг на новой строке)", True),
                ToolParameter("description", "string",
                              "Описание шаблона", False, ""),
            ],
            handler=tool_create_template,
            category="workflow",
        ),

        # ─── Part 10: Semantic Search V2 ────────────────────────────
        Tool(
            name="knowledge_add",
            description=(
                "Добавить знание в базу знаний. Знания индексируются "
                "для семантического поиска и могут быть найдены по смыслу."
            ),
            parameters=[
                ToolParameter("content", "string", "Содержимое знания", True),
                ToolParameter("category", "string",
                              "Категория: answer/document/conversation/fact/skill/business/general",
                              False, "general"),
                ToolParameter("source", "string", "Источник", False, ""),
                ToolParameter("tags", "string",
                              "Теги через запятую", False, ""),
            ],
            handler=tool_knowledge_add,
            category="knowledge",
        ),
        Tool(
            name="knowledge_search",
            description=(
                "Семантический поиск по базе знаний. "
                "Находит релевантные знания по смыслу, а не по точному совпадению."
            ),
            parameters=[
                ToolParameter("query", "string", "Поисковый запрос", True),
                ToolParameter("category", "string",
                              "Фильтр по категории", False, ""),
                ToolParameter("max_results", "number",
                              "Максимум результатов", False, 5),
            ],
            handler=tool_knowledge_search,
            category="knowledge",
        ),

        # ─── Part 10: Confidence Tracker ────────────────────────────
        Tool(
            name="confidence_check",
            description=(
                "Оценить уверенность в ответе. Показывает: уровень "
                "уверенности, факторы неопределённости, нужен ли "
                "дополнительный поиск."
            ),
            parameters=[
                ToolParameter("text", "string", "Текст для оценки", True),
                ToolParameter("source_count", "number",
                              "Количество источников", False, 1),
                ToolParameter("source_agreement", "number",
                              "Согласованность источников (0-1)", False, 0.5),
            ],
            handler=tool_confidence_check,
            category="confidence",
        ),

        # ─── Part 10: Adaptive Query Expansion ──────────────────────
        Tool(
            name="expand_query",
            description=(
                "Расширить/улучшить поисковый запрос. "
                "Добавляет синонимы, контекстные термины, временные маркеры. "
                "Помогает найти больше релевантных результатов."
            ),
            parameters=[
                ToolParameter("query", "string", "Исходный запрос", True),
                ToolParameter("context", "string",
                              "Контекст для расширения", False, ""),
                ToolParameter("strategy", "string",
                              "Стратегия: synonym/related/specific/broad/temporal/contextual",
                              False, "synonym"),
            ],
            handler=tool_expand_query,
            category="search",
        ),
        Tool(
            name="find_gaps",
            description=(
                "Найти пробелы в ответе: чего не хватает? "
                "Анализирует полноту, наличие данных, подтверждений."
            ),
            parameters=[
                ToolParameter("query", "string", "Исходный вопрос", True),
                ToolParameter("answer", "string", "Текущий ответ", True),
                ToolParameter("confidence", "number",
                              "Текущая уверенность (0-1)", False, 0.5),
            ],
            handler=tool_find_gaps,
            category="search",
        ),

        # ─── Part 10: Task Prioritizer ──────────────────────────────
        Tool(
            name="task_add",
            description=(
                "Добавить задачу в умную очередь с приоритетом. "
                "Задачи сортируются по приоритету, дедлайну, "
                "и возрасту (anti-starvation)."
            ),
            parameters=[
                ToolParameter("name", "string", "Название задачи", True),
                ToolParameter("priority", "string",
                              "Приоритет: critical/high/medium/low/background",
                              False, "medium"),
                ToolParameter("task_type", "string",
                              "Тип задачи: general/api/research/report",
                              False, "general"),
                ToolParameter("deadline_sec", "number",
                              "Дедлайн в секундах (0 = нет)", False, 0),
            ],
            handler=tool_task_add,
            category="tasks",
        ),
        Tool(
            name="task_queue",
            description=(
                "Показать очередь задач, план выполнения, "
                "оценку времени."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: list/plan/next/stats",
                              False, "list"),
            ],
            handler=tool_task_queue,
            category="tasks",
        ),

        # ─── Part 10: Context Compressor ────────────────────────────
        Tool(
            name="summarize_text",
            description=(
                "Суммаризировать текст (экстрактивная суммаризация). "
                "Выбирает ключевые предложения. Для длинных текстов "
                "используется рекурсивная суммаризация."
            ),
            parameters=[
                ToolParameter("text", "string",
                              "Текст для суммаризации", True),
                ToolParameter("ratio", "number",
                              "Степень сжатия (0.1-0.9, меньше = короче)",
                              False, 0.3),
                ToolParameter("recursive", "boolean",
                              "Рекурсивная суммаризация (для очень длинных)",
                              False, False),
            ],
            handler=tool_summarize_text,
            category="text",
        ),

        # ─── Part 10: Time & Relevance ──────────────────────────────
        Tool(
            name="check_freshness",
            description=(
                "Проверить актуальность данных. Извлекает даты, "
                "оценивает свежесть, даёт рекомендацию об обновлении. "
                "«Этот ответ основан на данных за 2023 год — проверить?»"
            ),
            parameters=[
                ToolParameter("text", "string", "Текст для проверки", True),
            ],
            handler=tool_check_freshness,
            category="analysis",
        ),
        Tool(
            name="time_decay",
            description=(
                "Применить временное затухание к оценке. "
                "Учитывает возраст данных для корректировки скора."
            ),
            parameters=[
                ToolParameter("score", "number", "Базовый скор (0-1)", True),
                ToolParameter("age_days", "number",
                              "Возраст данных в днях", True),
                ToolParameter("method", "string",
                              "Метод: exponential/linear/hyperbolic",
                              False, "exponential"),
            ],
            handler=tool_time_decay,
            category="analysis",
        ),

        # ─── Part 11: Integration Layer ─────────────────────────────
        Tool(
            name="run_chain",
            description=(
                "Запустить цепочку инструментов. Цепочки объединяют "
                "несколько tools в pipeline с передачей данных между шагами."
            ),
            parameters=[
                ToolParameter("chain_name", "string",
                              "Имя цепочки (research_summarize, confidence_check_search, "
                              "freshness_update, finance_report)", True),
                ToolParameter("query", "string",
                              "Входной запрос / данные", False, ""),
            ],
            handler=tool_run_chain,
            category="integration",
        ),
        Tool(
            name="tool_health",
            description=(
                "Показать здоровье инструментов: какие работают, "
                "какие деградируют, какие отключены circuit breaker."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: report/unhealthy/slow/stats",
                              False, "report"),
            ],
            handler=tool_health_check,
            category="integration",
        ),
        Tool(
            name="parallel_tools",
            description=(
                "Выполнить несколько инструментов параллельно. "
                "Принимает список вызовов и возвращает все результаты."
            ),
            parameters=[
                ToolParameter("calls", "string",
                              "Вызовы в формате: tool1:param1=val1;tool2:param2=val2",
                              True),
            ],
            handler=tool_parallel_execute,
            category="integration",
        ),
        Tool(
            name="list_chains",
            description=(
                "Показать все доступные цепочки инструментов."
            ),
            parameters=[],
            handler=tool_list_chains,
            category="integration",
        ),

        # ─── Part 12: Production Hardening ──────────────────────────────
        Tool(
            name="system_health",
            description=(
                "Полный системный отчёт: аптайм, здоровье подсистем, "
                "CPU, RAM, диск, активные алерты."
            ),
            parameters=[
                ToolParameter("section", "string",
                              "Секция: full/health/system/requests/errors/alerts",
                              False, "full"),
            ],
            handler=tool_system_health,
            category="production",
        ),
        Tool(
            name="rate_limit_info",
            description=(
                "Показать статус rate-лимитов: кто заблокирован, "
                "сколько запросов осталось."
            ),
            parameters=[
                ToolParameter("key", "string",
                              "Конкретный ключ (user_id или tool_name), "
                              "пусто = общая статистика",
                              False, ""),
            ],
            handler=tool_rate_limit_info,
            category="production",
        ),
        Tool(
            name="error_report",
            description=(
                "Отчёт об ошибках: последние ошибки, топ по частоте, "
                "статистика по типам."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: recent/top/stats/clear",
                              False, "recent"),
            ],
            handler=tool_error_report,
            category="production",
        ),
        Tool(
            name="uptime_info",
            description=(
                "Информация об аптайме системы: время работы, "
                "перезагрузки, простои."
            ),
            parameters=[],
            handler=tool_uptime_info,
            category="production",
        ),

        # ─── MESSAGING: WhatsApp, Telegram, Email ───────────────────
        Tool(
            name="send_whatsapp",
            description=(
                "Отправить сообщение в WhatsApp через Green-API. "
                "Укажи contact_name (имя контакта из базы) или phone (номер телефона). "
                "message — текст сообщения."
            ),
            parameters=[
                ToolParameter("contact_name", "string",
                              "Имя контакта (ищет в БД)", False),
                ToolParameter("phone", "string",
                              "Номер телефона (79001234567)", False),
                ToolParameter("message", "string",
                              "Текст сообщения", True),
            ],
            handler=tool_send_whatsapp,
            category="messaging",
            needs_db=True,
        ),
        Tool(
            name="send_telegram",
            description=(
                "Отправить личное сообщение в Telegram. "
                "Можно указать username (@user), contact_name (имя — 'Милана'), "
                "или chat_id. Имя разрешается автоматически через адресную книгу."
            ),
            parameters=[
                ToolParameter("username", "string",
                              "Telegram username (напр. @DurdyP)", False),
                ToolParameter("contact_name", "string",
                              "Имя контакта (напр. 'Милана', 'Кирилл')", False),
                ToolParameter("chat_id", "integer",
                              "Chat ID (числовой)", False),
                ToolParameter("message", "string",
                              "Текст сообщения", True),
            ],
            handler=tool_send_telegram,
            category="messaging",
            needs_db=True,
        ),
        Tool(
            name="read_telegram_chat",
            description=(
                "Прочитать историю сообщений Telegram чата. "
                "Работает по username, chat_id или contact_name (имя). "
                "Показывает последние сообщения — кто что писал."
            ),
            parameters=[
                ToolParameter("username", "string",
                              "Telegram username (напр. @milana_sagomonyan)", False),
                ToolParameter("contact_name", "string",
                              "Имя контакта (напр. 'Милана')", False),
                ToolParameter("chat_id", "integer",
                              "Chat ID (числовой)", False),
                ToolParameter("limit", "integer",
                              "Количество сообщений (макс 50)", False, "20"),
                ToolParameter("days", "integer",
                              "За сколько дней", False, "30"),
            ],
            handler=tool_read_telegram_chat,
            category="messaging",
            needs_db=True,
        ),
        # ── Contact Book (Smart Name Resolution) ──
        Tool(
            name="link_contact",
            description=(
                "Привязать контактные данные к имени. "
                "Используй когда пользователь говорит: 'запомни что у Миланы телеграм @milana', "
                "'email Кирилла — kirill@mail.ru', 'телефон мамы +79001234567'."
            ),
            parameters=[
                ToolParameter("name", "string",
                              "Имя контакта (напр. 'Милана', 'Кирилл', 'мама')", True),
                ToolParameter("telegram", "string",
                              "Telegram username (напр. @milana_sagomonyan)", False),
                ToolParameter("phone", "string",
                              "Номер телефона", False),
                ToolParameter("email", "string",
                              "Email адрес", False),
                ToolParameter("whatsapp", "string",
                              "WhatsApp номер", False),
            ],
            handler=tool_link_contact,
            category="contacts",
            needs_db=True,
        ),
        Tool(
            name="resolve_contact",
            description=(
                "Найти контакт по имени/нику/прозвищу. "
                "Поддерживает падежи ('Милане'), уменьшительные ('Серёга'→Сергей), "
                "ники ('макс'→Максим). Возвращает все привязанные данные."
            ),
            parameters=[
                ToolParameter("name", "string",
                              "Имя, ник или прозвище контакта", True),
            ],
            handler=tool_resolve_contact,
            category="contacts",
            needs_db=True,
        ),
        Tool(
            name="list_contacts",
            description="Показать все контакты из адресной книги с привязанными данными.",
            parameters=[],
            handler=tool_list_contacts,
            category="contacts",
        ),
        Tool(
            name="send_email",
            description=(
                "Отправить email через Gmail API. "
                "Укажи to (email адрес) или contact_name. "
                "subject — тема, body — текст письма."
            ),
            parameters=[
                ToolParameter("to", "string",
                              "Email получателя", False),
                ToolParameter("subject", "string",
                              "Тема письма", False, ""),
                ToolParameter("body", "string",
                              "Текст письма", True),
                ToolParameter("contact_name", "string",
                              "Имя контакта из БД (если нет to)", False),
            ],
            handler=tool_send_email,
            category="messaging",
            needs_db=True,
        ),
        Tool(
            name="get_emails",
            description=(
                "Получить непрочитанные письма из Gmail. "
                "Показывает отправителя, тему, дату, превью."
            ),
            parameters=[
                ToolParameter("account", "string",
                              "Аккаунт: work/personal/default (пусто = все)", False, ""),
                ToolParameter("max_results", "number",
                              "Максимум писем", False, 5),
            ],
            handler=tool_get_emails,
            category="messaging",
        ),

        # ─── FILES: Создание Excel, Word, PDF ───────────────────────
        Tool(
            name="create_file",
            description=(
                "Создать файл (Excel, Word, PDF, CSV, TXT, JSON) по описанию. "
                "DeepSeek определяет структуру и генерирует файл. "
                "Файл будет отправлен пользователю."
            ),
            parameters=[
                ToolParameter("description", "string",
                              "Описание файла (что создать)", True),
                ToolParameter("file_format", "string",
                              "Формат: xlsx/docx/pdf/csv/txt/json (авто если пусто)", False, ""),
            ],
            handler=tool_create_file,
            category="files",
        ),
        Tool(
            name="create_excel",
            description=(
                "Создать Excel таблицу с данными и отправить пользователю. "
                "Можно указать заголовки и строки, или пустые — будет примерочная."
            ),
            parameters=[
                ToolParameter("title", "string",
                              "Название таблицы", False, "Таблица"),
                ToolParameter("headers", "string",
                              "Заголовки через запятую: Имя,Возраст,Город", False, ""),
                ToolParameter("rows", "string",
                              "Строки через ; колонки через , : Анна,28,Москва;Иван,35,Питер", False, ""),
            ],
            handler=tool_create_excel,
            category="files",
        ),

        # ─── TELEGRAM BOT API: Chat Management ──────────────────────
        Tool(
            name="telegram_get_chat_info",
            description=(
                "Получить информацию о Telegram чате по chat_id. "
                "Показывает название, тип, bio, описание."
            ),
            parameters=[
                ToolParameter("chat_id", "number",
                              "ID чата (числовой)", True),
            ],
            handler=tool_telegram_get_chat_info,
            category="telegram",
        ),
        Tool(
            name="telegram_forward_message",
            description=(
                "Переслать сообщение между Telegram чатами. "
                "Бот должен быть участником обоих чатов."
            ),
            parameters=[
                ToolParameter("from_chat_id", "number",
                              "ID чата-источника", True),
                ToolParameter("to_chat_id", "number",
                              "ID чата-получателя", True),
                ToolParameter("message_id", "number",
                              "ID сообщения для пересылки", True),
            ],
            handler=tool_telegram_forward_message,
            category="telegram",
        ),
        Tool(
            name="telegram_pin_message",
            description=(
                "Закрепить сообщение в Telegram чате. "
                "Бот должен быть админом."
            ),
            parameters=[
                ToolParameter("chat_id", "number", "ID чата", True),
                ToolParameter("message_id", "number",
                              "ID сообщения для закрепления", True),
            ],
            handler=tool_telegram_pin_message,
            category="telegram",
        ),
        Tool(
            name="telegram_manage_chat",
            description=(
                "Управление Telegram чатом/группой через Bot API. "
                "Действия: info (инфо), set_title, set_description, "
                "ban/unban пользователя, get_members_count. "
                "Бот должен быть админом группы."
            ),
            parameters=[
                ToolParameter("action", "string",
                              "Действие: info/set_title/set_description/ban/unban/get_members_count",
                              False, "info"),
                ToolParameter("chat_id", "number", "ID чата/группы", True),
                ToolParameter("user_id", "number",
                              "ID пользователя (для ban/unban)", False),
                ToolParameter("title", "string",
                              "Новое название (для set_title)", False),
                ToolParameter("description", "string",
                              "Описание (для set_description)", False),
            ],
            handler=tool_telegram_manage_chat,
            category="telegram",
        ),
        Tool(
            name="telegram_send_photo",
            description=(
                "Отправить фото в Telegram чат. "
                "Нужен chat_id и путь к файлу."
            ),
            parameters=[
                ToolParameter("chat_id", "number", "ID чата", True),
                ToolParameter("photo_path", "string",
                              "Путь к файлу фото", True),
                ToolParameter("caption", "string",
                              "Подпись к фото", False, ""),
            ],
            handler=tool_telegram_send_photo,
            category="telegram",
        ),

        # ─── Sandbox / File Operations (Manus+ level) ──────────────
        Tool(
            name="sandbox_read_file",
            description=(
                "Прочитать файл (текст, Python, CSV, Excel, PDF). "
                "Поддержка диапазона строк."
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к файлу", True),
                ToolParameter("start_line", "number",
                              "Начальная строка (опц.)", False),
                ToolParameter("end_line", "number",
                              "Конечная строка (опц.)", False),
            ],
            handler=tool_sandbox_read_file,
            category="sandbox",
        ),
        Tool(
            name="sandbox_edit_file",
            description=(
                "Редактировать файл БЕЗ разрушения архитектуры. "
                "Точечные замены с бэкапом + валидация Python через AST. "
                "Edits: [{\"find\": \"old\", \"replace\": \"new\"}, "
                "{\"line\": 10, \"replace\": \"...\"}, "
                "{\"insert_after_line\": 5, \"content\": \"...\"}]"
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к файлу", True),
                ToolParameter("edits", "string",
                              "JSON массив правок: [{find,replace}, {line,replace}, {insert_after_line,content}]",
                              True),
                ToolParameter("create_backup", "boolean",
                              "Создать бэкап (по умолчанию да)", False, True),
            ],
            handler=tool_sandbox_edit_file,
            category="sandbox",
        ),
        Tool(
            name="sandbox_create_file",
            description="Создать новый файл с содержимым. Валидация .py через AST.",
            parameters=[
                ToolParameter("path", "string", "Путь к файлу", True),
                ToolParameter("content", "string", "Содержимое файла", True),
            ],
            handler=tool_sandbox_create_file,
            category="sandbox",
        ),
        Tool(
            name="sandbox_run_code",
            description=(
                "Выполнить Python код в безопасной песочнице. "
                "Ограничены опасные модули (subprocess, shutil, ctypes). "
                "Таймаут 30 сек."
            ),
            parameters=[
                ToolParameter("code", "string",
                              "Python код для выполнения", True),
                ToolParameter("timeout", "number",
                              "Таймаут в секундах", False, 30),
            ],
            handler=tool_sandbox_run_code,
            category="sandbox",
        ),
        Tool(
            name="sandbox_search_files",
            description="Поиск текста в файлах (grep). Поддержка regex.",
            parameters=[
                ToolParameter("pattern", "string", "Что искать", True),
                ToolParameter("directory", "string",
                              "Директория поиска", False),
                ToolParameter("extensions", "string",
                              "Расширения через запятую: .py,.txt,.csv", False),
                ToolParameter("regex", "boolean", "Regex поиск", False, False),
            ],
            handler=tool_sandbox_search_files,
            category="sandbox",
        ),
        Tool(
            name="sandbox_list_dir",
            description="Показать структуру директории (дерево файлов).",
            parameters=[
                ToolParameter("path", "string", "Путь к директории", False),
                ToolParameter("max_depth", "number",
                              "Глубина (по умолчанию 3)", False, 3),
            ],
            handler=tool_sandbox_list_dir,
            category="sandbox",
        ),
        Tool(
            name="sandbox_csv_read",
            description="Прочитать CSV/TSV файл с форматированием в таблицу.",
            parameters=[
                ToolParameter("path", "string", "Путь к CSV файлу", True),
                ToolParameter("max_rows", "number",
                              "Макс. строк для показа", False, 30),
            ],
            handler=tool_sandbox_csv_read,
            category="sandbox",
        ),
        Tool(
            name="sandbox_csv_edit",
            description=(
                "Редактировать CSV: add_row, edit_cell, delete_row, add_column, sort. "
                "Operations: [{\"op\": \"add_row\", \"data\": [\"v1\",\"v2\"]}, "
                "{\"op\": \"edit_cell\", \"row\": 0, \"col\": 1, \"value\": \"new\"}]"
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к CSV файлу", True),
                ToolParameter("operations", "string",
                              "JSON массив операций", True),
            ],
            handler=tool_sandbox_csv_edit,
            category="sandbox",
        ),

        # ─── Wide Research (UNIQUE — лучше Manus) ──────────────────
        Tool(
            name="wide_research",
            description=(
                "🔬 Широкое исследование: параллельные суб-агенты, "
                "детекция противоречий, скоринг уверенности. "
                "Лучше Manus Wide Research."
            ),
            parameters=[
                ToolParameter("query", "string", "Тема исследования", True),
                ToolParameter("max_sources", "number",
                              "Макс. источников на суб-запрос", False, 5),
            ],
            handler=tool_wide_research,
            category="research",
        ),
        Tool(
            name="quick_research",
            description="Быстрое исследование без LLM — только поиск + анализ.",
            parameters=[
                ToolParameter("query", "string", "Запрос", True),
                ToolParameter("max_sources", "number",
                              "Макс. источников", False, 3),
            ],
            handler=tool_quick_research,
            category="research",
        ),
        Tool(
            name="compare_research",
            description=(
                "🔥 УНИКАЛЬНО: сравнить N объектов по M критериям. "
                "Параллельное исследование каждого. "
                "items через запятую, criteria через запятую."
            ),
            parameters=[
                ToolParameter("items", "string",
                              "Объекты для сравнения через запятую", True),
                ToolParameter("criteria", "string",
                              "Критерии через запятую (опц.)", False),
            ],
            handler=tool_compare_research,
            category="research",
        ),

        # ─── Data Analysis (встроенный движок) ─────────────────────
        Tool(
            name="analyze_data",
            description=(
                "📊 Полный EDA (Exploratory Data Analysis) файла. "
                "CSV/Excel/JSON → статистика, корреляции, графики."
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к файлу данных", True),
                ToolParameter("generate_charts", "boolean",
                              "Генерировать графики", False, True),
            ],
            handler=tool_analyze_data,
            category="data",
        ),
        Tool(
            name="create_chart",
            description=(
                "📈 Создать график (bar, line, pie, scatter) из данных файла."
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к файлу данных", True),
                ToolParameter("x_column", "string", "Столбец X", True),
                ToolParameter("y_column", "string", "Столбец Y", True),
                ToolParameter("chart_type", "string",
                              "Тип: bar, line, pie, scatter", False, "bar"),
                ToolParameter("title", "string", "Заголовок графика", False),
            ],
            handler=tool_create_chart,
            category="data",
        ),
        Tool(
            name="data_filter",
            description=(
                "🔍 Фильтрация данных по условию: equals, contains, "
                "greater, less, not_empty, starts_with."
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к файлу данных", True),
                ToolParameter("column", "string", "Столбец", True),
                ToolParameter("condition", "string",
                              "Условие: equals/contains/greater/less/not_empty/starts_with",
                              True),
                ToolParameter("value", "string", "Значение", False, ""),
            ],
            handler=tool_data_filter,
            category="data",
        ),
        Tool(
            name="data_group_by",
            description=(
                "📊 Группировка данных с агрегацией: count, sum, avg, min, max."
            ),
            parameters=[
                ToolParameter("path", "string", "Путь к файлу данных", True),
                ToolParameter("group_column", "string",
                              "Столбец группировки", True),
                ToolParameter("agg_column", "string",
                              "Столбец агрегации (опц.)", False),
                ToolParameter("agg_func", "string",
                              "Функция: count/sum/avg/min/max", False, "count"),
            ],
            handler=tool_data_group_by,
            category="data",
        ),
        Tool(
            name="data_stats",
            description="📊 Детальная статистика по столбцу или всем числовым столбцам.",
            parameters=[
                ToolParameter("path", "string", "Путь к файлу данных", True),
                ToolParameter("column", "string", "Столбец (опц.)", False),
            ],
            handler=tool_data_stats,
            category="data",
        ),
        # ── v6: Persona & Proactive ──
        Tool(
            name="persona_stats",
            description="🧠 Статистика персоны: сколько пользователей изучено, группы сходства.",
            parameters=[],
            handler=tool_persona_stats,
            category="persona",
        ),
        Tool(
            name="persona_retrain",
            description="🔄 Принудительный retrain персоны из истории чатов.",
            parameters=[
                ToolParameter("days", "integer",
                              "За сколько дней перечитать историю", False, "3"),
            ],
            handler=tool_persona_retrain,
            category="persona",
        ),
        Tool(
            name="persona_style",
            description="📋 Показать стиль-гайд пользователя (как он общается).",
            parameters=[
                ToolParameter("chat_id", "integer",
                              "Chat ID пользователя (0 = владелец)", False, "0"),
            ],
            handler=tool_persona_style,
            category="persona",
        ),
        Tool(
            name="proactive_status",
            description="⚡ Статус проактивного движка: задачи, аномалии, фильтры.",
            parameters=[],
            handler=tool_proactive_status,
            category="proactive",
        ),
        Tool(
            name="add_important_keyword",
            description="🔑 Добавить ключевое слово для проактивного мониторинга сообщений.",
            parameters=[
                ToolParameter("keyword", "string", "Ключевое слово", True),
                ToolParameter("weight", "number",
                              "Вес важности (0-2)", False, "1.0"),
            ],
            handler=tool_add_important_keyword,
            category="proactive",
        ),
    ]

    for tool in tools:
        tool_registry.register(tool)

    logger.info(f"Зарегистрировано {len(tools)} бизнес-инструментов агента")
    return len(tools)


# ═══════════════════════════════════════════════════════════════════════════════
# PART 8: PLUGIN TOOLS (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_plugin_connect(
    name: str,
    base_url: str,
    api_key: str = "",
    plugin_type: str = "REST_API",
    **kwargs,
) -> ToolResult:
    """Подключить внешний API как плагин."""
    from pds_ultimate.core.plugin_system import PluginType, plugin_manager

    try:
        # Маппинг строки в enum
        type_map = {t.value: t for t in PluginType}
        p_type = type_map.get(plugin_type.upper(), PluginType.REST_API)

        plugin = await plugin_manager.register_plugin(
            name=name,
            base_url=base_url,
            api_key=api_key if api_key else None,
            plugin_type=p_type,
            user_id=kwargs.get("_user_id", "system"),
        )

        return ToolResult(
            "plugin_connect", True,
            f"✅ Плагин «{plugin.name}» подключён\n"
            f"  🔗 URL: {plugin.base_url}\n"
            f"  📋 Тип: {plugin.plugin_type.value}\n"
            f"  🆔 ID: {plugin.id}",
            data={"plugin_id": plugin.id, "name": plugin.name},
        )
    except Exception as e:
        return ToolResult(
            "plugin_connect", False, "",
            error=f"Ошибка подключения плагина: {e}",
        )


async def tool_plugin_execute(
    plugin_name: str,
    endpoint: str,
    method: str = "GET",
    body: str = "",
    **kwargs,
) -> ToolResult:
    """Выполнить запрос через плагин."""
    import json as _json

    from pds_ultimate.core.plugin_system import plugin_manager

    try:
        plugin = plugin_manager.get_by_name(plugin_name)
        if not plugin:
            return ToolResult(
                "plugin_execute", False, "",
                error=f"Плагин «{plugin_name}» не найден",
            )

        # Парсим тело запроса
        json_body = None
        if body:
            try:
                json_body = _json.loads(body)
            except _json.JSONDecodeError:
                json_body = {"data": body}

        result = await plugin_manager.execute(
            plugin_id=plugin.id,
            endpoint=endpoint,
            method=method.upper(),
            json_data=json_body,
        )

        # Форматируем ответ
        if isinstance(result, dict):
            output = _json.dumps(result, ensure_ascii=False, indent=2)[:3000]
        else:
            output = str(result)[:3000]

        return ToolResult(
            "plugin_execute", True,
            f"📡 {plugin_name} → {method.upper()} {endpoint}\n\n{output}",
            data=result if isinstance(result, dict) else {"response": output},
        )
    except Exception as e:
        return ToolResult(
            "plugin_execute", False, "",
            error=f"Ошибка вызова плагина: {e}",
        )


async def tool_plugin_list(**kwargs) -> ToolResult:
    """Список подключённых плагинов."""
    from pds_ultimate.core.plugin_system import plugin_manager

    stats = plugin_manager.get_stats()
    plugins = plugin_manager.get_active_plugins()

    if not plugins:
        return ToolResult(
            "plugin_list", True,
            "📋 Нет подключённых плагинов.\n"
            "Используй plugin_connect для подключения API.",
        )

    lines = [f"📋 Плагины ({stats['total']}):"]
    for p in plugins:
        lines.append(
            f"  • {p.name} [{p.plugin_type.value}] — {p.status.value}\n"
            f"    🔗 {p.base_url}"
        )

    return ToolResult(
        "plugin_list", True, "\n".join(lines),
        data=stats,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 8: AUTONOMY TOOLS (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_autonomous_task(
    goal: str,
    priority: str = "normal",
    deadline_hours: float = 0,
    **kwargs,
) -> ToolResult:
    """Создать автономную задачу."""
    from pds_ultimate.core.autonomy_engine import TaskPriority, autonomy_engine

    try:
        # Маппинг строки в приоритет
        priority_map = {
            "critical": TaskPriority.CRITICAL,
            "high": TaskPriority.HIGH,
            "normal": TaskPriority.MEDIUM,
            "medium": TaskPriority.MEDIUM,
            "low": TaskPriority.LOW,
            "background": TaskPriority.BACKGROUND,
        }
        p = priority_map.get(priority.lower(), TaskPriority.MEDIUM)

        # Дедлайн
        from datetime import datetime, timedelta
        deadline = None
        if deadline_hours and float(deadline_hours) > 0:
            deadline = datetime.utcnow() + timedelta(hours=float(deadline_hours))

        task = autonomy_engine.create_task(
            title=goal,
            description=goal,
            priority=p,
            deadline=deadline,
            owner_id=kwargs.get("_user_id", 0),
            chat_id=kwargs.get("_chat_id", 0),
        )

        lines = [
            "🤖 Автономная задача создана:",
            f"  🆔 ID: {task.id}",
            f"  🎯 Цель: {task.title}",
            f"  ⚡ Приоритет: {priority}",
        ]
        if deadline:
            lines.append(f"  ⏰ Дедлайн: {deadline.strftime('%Y-%m-%d %H:%M')}")

        return ToolResult(
            "autonomous_task", True, "\n".join(lines),
            data={"task_id": task.id, "status": task.status.value},
        )
    except Exception as e:
        return ToolResult(
            "autonomous_task", False, "",
            error=f"Ошибка создания задачи: {e}",
        )


async def tool_task_status(task_id: str = "", **kwargs) -> ToolResult:
    """Статус автономных задач."""
    from pds_ultimate.core.autonomy_engine import autonomy_engine

    try:
        if task_id:
            task = autonomy_engine.get_task(task_id)
            if not task:
                return ToolResult(
                    "task_status", False, "",
                    error=f"Задача {task_id} не найдена",
                )
            lines = [
                f"📋 Задача {task.id}:",
                f"  🎯 {task.title}",
                f"  📊 Статус: {task.status.value}",
                f"  📈 Прогресс: {task.progress:.0%}",
                f"  🔧 Шагов: {len(task.steps)}",
            ]
            if task.corrections:
                lines.append(f"  🔄 Коррекций: {len(task.corrections)}")
            return ToolResult(
                "task_status", True, "\n".join(lines),
                data={"task_id": task.id, "status": task.status.value,
                      "progress": task.progress},
            )

        # Все активные
        stats = autonomy_engine.get_stats()
        queue = autonomy_engine.format_queue()
        return ToolResult(
            "task_status", True,
            f"📋 Автономные задачи:\n{queue}\n\n"
            f"📊 Всего: {stats['total']}, Активных: {stats['active']}",
            data=stats,
        )
    except Exception as e:
        return ToolResult(
            "task_status", False, "",
            error=f"Ошибка получения статуса: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 8: MEMORY V2 TOOLS (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_learn_skill(
    name: str,
    pattern: str,
    strategy: str,
    **kwargs,
) -> ToolResult:
    """Научить агента новому навыку."""
    from pds_ultimate.core.memory_v2 import memory_v2

    try:
        skill = memory_v2.learn_skill(
            name=name,
            pattern=pattern,
            strategy=strategy,
        )
        return ToolResult(
            "learn_skill", True,
            f"🎓 Навык «{skill.name}» сохранён!\n"
            f"  📋 Паттерн: {pattern}\n"
            f"  💡 Стратегия: {strategy}",
            data=skill.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "learn_skill", False, "",
            error=f"Ошибка сохранения навыка: {e}",
        )


async def tool_memory_stats(**kwargs) -> ToolResult:
    """Статистика памяти v2."""
    from pds_ultimate.core.memory_v2 import memory_v2

    try:
        stats = memory_v2.get_stats()

        lines = [
            "🧠 Статистика памяти v2:",
            f"  🎓 Навыков: {stats['skills']}",
            f"  ⚠️ Ошибок записано: {stats['failures']}",
            f"  📊 Паттернов: {stats['patterns']}",
        ]

        if stats.get("top_skills"):
            lines.append("\n🏆 Топ навыки:")
            for s in stats["top_skills"]:
                lines.append(f"  • {s['name']} ({s['success_rate']})")

        fail_stats = stats.get("failure_stats", {})
        if fail_stats.get("by_type"):
            lines.append("\n📊 Ошибки по типу:")
            for t, c in fail_stats["by_type"].items():
                lines.append(f"  • {t}: {c}")

        return ToolResult(
            "memory_stats", True, "\n".join(lines),
            data=stats,
        )
    except Exception as e:
        return ToolResult(
            "memory_stats", False, "",
            error=f"Ошибка получения статистики: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 9: SMART TRIGGERS (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_set_trigger(
    name: str,
    trigger_type: str = "threshold",
    field: str = "",
    operator: str = ">",
    value: str = "",
    severity: str = "warning",
    template: str = "",
    **kwargs,
) -> ToolResult:
    """Установить умный триггер."""
    from pds_ultimate.core.smart_triggers import (
        ComparisonOp,
        TriggerCondition,
        trigger_manager,
    )

    try:
        # Если указан шаблон — используем его
        if template:
            template_kwargs = {}
            if value:
                # Парсим значение для шаблона
                try:
                    template_kwargs["threshold"] = float(value)
                except ValueError:
                    template_kwargs["supplier_name"] = value

            trigger = trigger_manager.create_from_template(
                template, **template_kwargs,
            )
            trigger.name = name or trigger.name
        else:
            # Создаём пользовательский триггер
            condition = None
            if field and value:
                try:
                    op = ComparisonOp(operator)
                except ValueError:
                    op = ComparisonOp.GT

                try:
                    val = float(value)
                except ValueError:
                    val = value

                condition = TriggerCondition(
                    field=field,
                    operator=op,
                    value=val,
                )

            trigger = trigger_manager.create_trigger(
                name=name,
                trigger_type=trigger_type,
                condition=condition,
                severity=severity,
            )

        return ToolResult(
            "set_trigger", True,
            f"🔔 Триггер «{trigger.name}» создан!\n"
            f"  🆔 ID: {trigger.id}\n"
            f"  📋 Тип: {trigger.trigger_type.value}\n"
            f"  ⚡ Серьёзность: {trigger.severity.value}\n"
            f"  📌 Условие: {trigger.condition.describe() if trigger.condition else 'custom'}",
            data=trigger.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "set_trigger", False, "",
            error=f"Ошибка создания триггера: {e}",
        )


async def tool_list_triggers(
    show_history: bool = False,
    **kwargs,
) -> ToolResult:
    """Список триггеров и алертов."""
    from pds_ultimate.core.smart_triggers import trigger_manager

    try:
        triggers_text = trigger_manager.format_triggers_list()
        stats = trigger_manager.get_stats()

        lines = [triggers_text]
        lines.append(
            f"\n📊 Всего: {stats['total']}, "
            f"активных: {stats['active']}, "
            f"срабатываний: {stats['total_fires']}"
        )

        if show_history:
            recent = trigger_manager.history.get_recent(10)
            if recent:
                lines.append("\n📜 Последние алерты:")
                for a in recent:
                    lines.append(f"  • {a.format_message()}")
            else:
                lines.append("\n📜 Алертов пока нет.")

        return ToolResult(
            "list_triggers", True, "\n".join(lines),
            data=stats,
        )
    except Exception as e:
        return ToolResult(
            "list_triggers", False, "",
            error=f"Ошибка получения триггеров: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 9: ANALYTICS DASHBOARD (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_dashboard(
    action: str = "show",
    metric_name: str = "",
    value: float = 0.0,
    unit: str = "",
    **kwargs,
) -> ToolResult:
    """Бизнес-дашборд."""
    from pds_ultimate.core.analytics_dashboard import analytics_dashboard

    try:
        if action == "record" and metric_name:
            analytics_dashboard.record_metric(
                name=metric_name,
                value=float(value),
                unit=unit,
            )
            return ToolResult(
                "dashboard", True,
                f"📊 Записано: {metric_name} = {value} {unit}",
            )
        elif action == "trend" and metric_name:
            report = analytics_dashboard.generate_trend_report()
            return ToolResult(
                "dashboard", True, report,
                data=analytics_dashboard.get_stats(),
            )
        elif action == "forecast" and metric_name:
            forecast = analytics_dashboard.forecast(metric_name)
            return ToolResult(
                "dashboard", True,
                f"📈 Прогноз {metric_name}: {forecast}",
                data={"forecast": forecast},
            )
        else:
            dashboard = analytics_dashboard.generate_dashboard()
            return ToolResult(
                "dashboard", True, dashboard,
                data=analytics_dashboard.get_stats(),
            )
    except Exception as e:
        return ToolResult(
            "dashboard", False, "",
            error=f"Ошибка дашборда: {e}",
        )


async def tool_kpi_track(
    action: str = "board",
    name: str = "",
    target: float = 0.0,
    value: float = 0.0,
    unit: str = "",
    **kwargs,
) -> ToolResult:
    """Отслеживание KPI."""
    from pds_ultimate.core.analytics_dashboard import analytics_dashboard

    try:
        if action == "create" and name:
            kpi = analytics_dashboard.create_kpi(
                name=name,
                target=float(target),
                unit=unit,
            )
            return ToolResult(
                "kpi_track", True,
                f"🎯 KPI «{kpi.name}» создан!\n"
                f"  📊 Цель: {kpi.target_value} {kpi.unit}\n"
                f"  📈 Прогресс: {kpi.progress_percent}%",
                data=kpi.to_dict(),
            )
        elif action == "update" and name:
            kpi = analytics_dashboard.update_kpi(name, float(value))
            if not kpi:
                return ToolResult(
                    "kpi_track", False, "",
                    error=f"KPI «{name}» не найден",
                )
            return ToolResult(
                "kpi_track", True,
                f"📊 KPI «{kpi.name}» обновлён!\n"
                f"  📈 {kpi.current_value:.0f}/{kpi.target_value:.0f} "
                f"{kpi.unit} [{kpi.progress_percent}%]\n"
                f"  📋 Статус: {kpi.status.value}",
                data=kpi.to_dict(),
            )
        else:
            board = analytics_dashboard.kpi_tracker.format_kpi_board()
            stats = analytics_dashboard.kpi_tracker.get_stats()
            return ToolResult(
                "kpi_track", True, board,
                data=stats,
            )
    except Exception as e:
        return ToolResult(
            "kpi_track", False, "",
            error=f"Ошибка KPI: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 9: CRM (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_rate_contact(
    name: str,
    rating: float,
    comment: str = "",
    category: str = "",
    **kwargs,
) -> ToolResult:
    """Оценить контакт/поставщика."""
    from pds_ultimate.core.crm_engine import crm_engine

    try:
        rating = max(1.0, min(5.0, float(rating)))

        if category:
            # Оценка поставщика по категории
            scorecard = crm_engine.rate_supplier(name, category, rating)
            if not scorecard:
                # Автосоздание контакта
                contact = crm_engine.add_contact(
                    name=name, contact_type="supplier",
                    rating=rating,
                )
                scorecard = crm_engine.rate_supplier(name, category, rating)

            return ToolResult(
                "rate_contact", True,
                f"📊 Оценка «{name}» [{category}]: {rating}/5\n"
                f"  🏆 Общий балл: {scorecard.overall_score}/5.0"
                if scorecard else f"⚠️ Не удалось оценить {name}",
                data=scorecard.to_dict() if scorecard else {},
            )
        else:
            # Общая оценка контакта
            contact = crm_engine.rate_contact(name, rating, comment)
            if not contact:
                # Автосоздание
                contact = crm_engine.add_contact(
                    name=name, rating=rating,
                )

            return ToolResult(
                "rate_contact", True,
                f"⭐ «{name}» оценён: {contact.star_rating} ({contact.rating}/5)"
                + (f"\n  💬 {comment}" if comment else ""),
                data=contact.to_dict(),
            )
    except Exception as e:
        return ToolResult(
            "rate_contact", False, "",
            error=f"Ошибка оценки: {e}",
        )


async def tool_crm_search(
    query: str = "",
    action: str = "search",
    contact_type: str = "",
    min_rating: float = 0.0,
    **kwargs,
) -> ToolResult:
    """Поиск в CRM."""
    from pds_ultimate.core.crm_engine import crm_engine

    try:
        if action == "pipeline":
            text = crm_engine.pipeline.format_pipeline()
            stats = crm_engine.pipeline.get_stats()
            return ToolResult(
                "crm_search", True, text,
                data=stats,
            )
        elif action == "stats":
            stats = crm_engine.get_stats()
            lines = [
                "📊 CRM Статистика:",
                f"  👤 Контактов: {stats['contacts']['total']}",
                f"  📊 Средний рейтинг: {stats['contacts']['avg_rating']}",
                f"  💼 Сделок: {stats['pipeline']['total']}",
                f"  💬 Взаимодействий: {stats['interactions']}",
                f"  📞 Ожидают follow-up: {stats['pending_followups']}",
            ]
            return ToolResult(
                "crm_search", True, "\n".join(lines),
                data=stats,
            )
        elif action == "add_contact" and query:
            contact = crm_engine.add_contact(
                name=query, contact_type=contact_type or "other",
            )
            return ToolResult(
                "crm_search", True,
                f"✅ Контакт «{contact.name}» добавлен (ID: {contact.id})",
                data=contact.to_dict(),
            )
        elif action == "add_deal" and query:
            deal = crm_engine.create_deal(title=query)
            return ToolResult(
                "crm_search", True,
                f"✅ Сделка «{deal.title}» создана (ID: {deal.id})",
                data=deal.to_dict(),
            )
        else:
            # Search
            contacts = crm_engine.search_contacts(
                query=query,
                contact_type=contact_type,
                min_rating=float(min_rating),
            )
            if not contacts:
                return ToolResult(
                    "crm_search", True,
                    f"🔍 По запросу «{query}» контактов не найдено.",
                )

            lines = [f"🔍 Найдено контактов: {len(contacts)}"]
            for c in contacts[:10]:
                lines.append(f"\n{c.format_card()}")
            return ToolResult(
                "crm_search", True, "\n".join(lines),
                data={"count": len(contacts),
                      "contacts": [c.to_dict() for c in contacts[:10]]},
            )
    except Exception as e:
        return ToolResult(
            "crm_search", False, "",
            error=f"Ошибка CRM: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 9: EVENING DIGEST (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_evening_digest(
    format: str = "full",
    revenue: float = 0.0,
    expenses: float = 0.0,
    orders_created: int = 0,
    tasks_completed: int = 0,
    **kwargs,
) -> ToolResult:
    """Вечерний дайджест."""
    from pds_ultimate.core.evening_digest import DaySummary, evening_digest

    try:
        summary = DaySummary(
            revenue=float(revenue),
            expenses=float(expenses),
            profit=float(revenue) - float(expenses),
            orders_created=int(orders_created),
            tasks_completed=int(tasks_completed),
        )
        evening_digest.record_day_summary(summary)

        if format == "short":
            text = evening_digest.generate_short_digest(summary)
        else:
            text = evening_digest.generate_digest(summary)

        return ToolResult(
            "evening_digest", True, text,
            data=summary.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "evening_digest", False, "",
            error=f"Ошибка дайджеста: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 9: WORKFLOW & TEMPLATES (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_create_template(
    name: str,
    template_type: str = "checklist",
    content: str = "",
    description: str = "",
    **kwargs,
) -> ToolResult:
    """Создать шаблон или чек-лист."""
    from pds_ultimate.core.workflow_engine import workflow_engine

    try:
        if template_type == "checklist" and content:
            # Создаём чек-лист из содержимого
            steps = [
                s.strip().lstrip("0123456789.-) ")
                for s in content.split("\n")
                if s.strip()
            ]
            checklist = workflow_engine.create_checklist(
                name=name,
                steps=steps,
                description=description,
            )
            return ToolResult(
                "create_template", True,
                f"📋 Чек-лист «{checklist.name}» создан!\n"
                f"{checklist.format_text()}",
                data=checklist.to_dict(),
            )
        else:
            # Создаём шаблон
            template = workflow_engine.create_template(
                name=name,
                template_type=template_type,
                content=content,
                description=description,
            )
            return ToolResult(
                "create_template", True,
                f"📝 Шаблон «{template.name}» создан!\n"
                f"  📋 Тип: {template.template_type.value}\n"
                f"  🆔 ID: {template.id}",
                data=template.to_dict(),
            )
    except Exception as e:
        return ToolResult(
            "create_template", False, "",
            error=f"Ошибка создания шаблона: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# BROWSER TOOLS (handlers) — Manus-level browsing
# ═══════════════════════════════════════════════════════════════════════════════

async def tool_web_search(query: str, max_results: int = 10, **kwargs) -> ToolResult:
    """
    Поиск в интернете — Manus-level.
    Primary: HttpxBrowser (всегда работает)
    Fallback: Playwright BrowserEngine (если установлен)
    """
    max_results = min(int(max_results), 20)

    from pds_ultimate.core.httpx_browser import httpx_browser

    try:
        results = await httpx_browser.search(query, max_results=max_results)

        if not results:
            return ToolResult(
                "web_search", True,
                f"По запросу «{query}» ничего не найдено.",
                data={"results": []},
            )

        lines = [f"🔍 Результаты поиска: «{query}» ({len(results)} шт.)\n"]
        for r in results:
            lines.append(f"  {r.position}. {r.title}")
            lines.append(f"     🔗 {r.url}")
            if r.snippet:
                lines.append(f"     {r.snippet[:150]}")
            lines.append("")

        return ToolResult(
            "web_search", True, "\n".join(lines),
            data={"results": [
                {"title": r.title, "url": r.url, "snippet": r.snippet}
                for r in results
            ]},
        )
    except Exception as e:
        return ToolResult("web_search", False, "",
                          error=f"Ошибка поиска: {e}")


async def tool_open_page(url: str, **kwargs) -> ToolResult:
    """
    Открыть страницу и извлечь данные — Manus-level.
    Primary: HttpxBrowser (всегда)
    Fallback: Playwright BrowserEngine
    """
    from pds_ultimate.core.httpx_browser import httpx_browser

    try:
        page = await httpx_browser.open_page(url)

        if not page.success:
            return ToolResult("open_page", False, "",
                              error=f"Не удалось загрузить: {url} ({page.error})")

        text = page.text[:4000] if page.text else ""
        extra = ""
        if page.text and len(page.text) > 4000:
            extra = f"\n\n... (ещё {len(page.text) - 4000} символов)"

        lines = [f"📄 {page.title}", f"🔗 {page.url}"]
        if page.headings:
            lines.append(f"📑 Заголовков: {len(page.headings)}")
        if page.tables:
            lines.append(f"📊 Таблиц: {len(page.tables)}")
        if page.links:
            lines.append(f"🔗 Ссылок: {len(page.links)}")
        lines.append(f"⏱ {page.load_time_ms}ms")
        lines.append("")
        if text:
            lines.append(text + extra)

        return ToolResult(
            "open_page", True, "\n".join(lines),
            data={
                "title": page.title,
                "url": page.url,
                "text_len": len(page.text or ""),
                "tables": page.tables[:3],
                "headings": page.headings[:10],
                "links_count": len(page.links),
            },
        )
    except Exception as e:
        return ToolResult("open_page", False, "",
                          error=f"Ошибка загрузки страницы: {e}")


async def tool_search_and_read(
    query: str,
    max_pages: int = 3,
    **kwargs,
) -> ToolResult:
    """
    Manus-level: Поиск + автоматическое чтение топ-N страниц.
    Одним вызовом: ищет → открывает → извлекает → возвращает.
    """
    max_pages = min(int(max_pages), 5)

    from pds_ultimate.core.httpx_browser import httpx_browser

    try:
        pages = await httpx_browser.search_and_extract(
            query, max_pages=max_pages, max_text_per_page=2500
        )

        if not pages:
            return ToolResult(
                "search_and_read", True,
                f"По запросу «{query}» не удалось получить контент со страниц.",
                data={"pages": []},
            )

        lines = [f"🔍📖 Поиск + чтение: «{query}» ({len(pages)} источников)\n"]

        for i, page in enumerate(pages, 1):
            lines.append(f"━━━ Источник {i}: {page.title} ━━━")
            lines.append(f"🔗 {page.url}")
            if page.text:
                lines.append(page.text[:2500])
            if page.tables:
                lines.append(f"\n📊 Таблиц: {len(page.tables)}")
                for tbl in page.tables[:2]:
                    for row in tbl[:5]:
                        lines.append("  | " + " | ".join(row))
            lines.append("")

        stats = httpx_browser.get_session_stats()
        lines.append(
            f"\n📊 Статистика: {stats['total_requests']} запросов, "
            f"{stats['total_bytes'] // 1024}KB, {stats['duration_ms']}ms"
        )

        return ToolResult(
            "search_and_read", True, "\n".join(lines),
            data={
                "pages": [
                    {"url": p.url, "title": p.title,
                     "text_len": len(p.text or "")}
                    for p in pages
                ],
            },
        )
    except Exception as e:
        return ToolResult("search_and_read", False, "",
                          error=f"Ошибка поиска и чтения: {e}")


async def tool_deep_web_research(
    query: str,
    max_sources: int = 5,
    **kwargs,
) -> ToolResult:
    """
    Manus-level глубокое исследование:
    Поиск → чтение страниц → переход по ссылкам → сбор всех данных.
    """
    max_sources = min(int(max_sources), 10)

    from pds_ultimate.core.httpx_browser import httpx_browser

    try:
        result = await httpx_browser.deep_search(
            query,
            max_sources=max_sources,
            follow_depth=1,
            max_text_per_page=2000,
        )

        findings = result.get("findings", [])
        if not findings:
            return ToolResult(
                "deep_web_research", True,
                f"По теме «{query}» не удалось собрать данные.",
                data=result,
            )

        lines = [
            f"🔬 Глубокое исследование: «{query}»",
            f"📖 Источников: {result['sources_count']}",
            f"📄 Страниц обработано: {result['pages_fetched']}",
            f"⏱ {result['duration_ms']}ms\n",
        ]

        for i, f in enumerate(findings, 1):
            lines.append(f"━━━ [{i}] {f['title']} ━━━")
            lines.append(f"🔗 {f['url']}")
            lines.append(f['text'][:2000])
            if f.get('tables'):
                lines.append(f"📊 Таблиц: {len(f['tables'])}")
            lines.append("")

        return ToolResult(
            "deep_web_research", True, "\n".join(lines),
            data=result,
        )
    except Exception as e:
        return ToolResult("deep_web_research", False, "",
                          error=f"Ошибка глубокого исследования: {e}")


async def tool_extract_page_data(
    url: str,
    focus: str = "",
    **kwargs,
) -> ToolResult:
    """Извлечь структурированные данные со страницы."""
    from pds_ultimate.core.httpx_browser import httpx_browser

    try:
        result = await httpx_browser.extract_structured(url, focus=focus)

        if result.get("error"):
            return ToolResult("extract_page_data", False, "",
                              error=result["error"])

        lines = [f"📄 {result.get('title', url)}"]
        lines.append(f"🔗 {url}")

        if result.get("meta"):
            desc = result["meta"].get("description", "")
            if desc:
                lines.append(f"📝 {desc[:200]}")

        if result.get("headings"):
            lines.append("\n📑 Заголовки:")
            for h in result["headings"][:15]:
                indent = "  " * int(h["level"][1])
                lines.append(f"{indent}• {h['text']}")

        if result.get("focused_text"):
            lines.append(f"\n🎯 Текст по теме «{focus}»:")
            lines.append(result["focused_text"][:3000])
        elif result.get("text"):
            lines.append("\n📄 Текст:")
            lines.append(result["text"][:3000])

        if result.get("tables"):
            lines.append(f"\n📊 Таблицы ({result['tables_count']}):")
            for tbl in result["tables"][:3]:
                for row in tbl[:8]:
                    lines.append("  | " + " | ".join(str(c) for c in row))
                lines.append("  ---")

        lines.append(
            f"\n📊 Ссылок: {result.get('links_count', 0)}, "
            f"Изображений: {result.get('images_count', 0)}"
        )

        return ToolResult(
            "extract_page_data", True, "\n".join(lines),
            data=result,
        )
    except Exception as e:
        return ToolResult("extract_page_data", False, "",
                          error=f"Ошибка извлечения данных: {e}")


async def tool_browser_screenshot(full_page: bool = False, **kwargs) -> ToolResult:
    """Скриншот текущей страницы."""
    from pds_ultimate.core.browser_engine import browser_engine

    try:
        path = await browser_engine.screenshot(full_page=bool(full_page))
        return ToolResult(
            "browser_screenshot", True,
            f"📸 Скриншот сохранён: {path}",
            data={"path": str(path)},
        )
    except RuntimeError as e:
        return ToolResult("browser_screenshot", False, "", error=str(e))
    except Exception as e:
        return ToolResult("browser_screenshot", False, "",
                          error=f"Ошибка скриншота: {e}")


async def tool_browser_click(selector: str, **kwargs) -> ToolResult:
    """Кликнуть по элементу."""
    from pds_ultimate.core.browser_engine import browser_engine

    try:
        await browser_engine.click(selector, human_like=True)
        # Ждём загрузку после клика
        await asyncio.sleep(1.0)
        info = await browser_engine.get_page_info()
        return ToolResult(
            "browser_click", True,
            f"✅ Кликнул по '{selector}'. Текущая страница: {info.title}",
            data={"url": info.url, "title": info.title},
        )
    except RuntimeError as e:
        return ToolResult("browser_click", False, "", error=str(e))
    except Exception as e:
        return ToolResult("browser_click", False, "",
                          error=f"Ошибка клика: {e}")


async def tool_browser_fill(selector: str, value: str, **kwargs) -> ToolResult:
    """Заполнить поле."""
    from pds_ultimate.core.browser_engine import browser_engine

    try:
        await browser_engine.fill(selector, value, human_like=True)
        return ToolResult(
            "browser_fill", True,
            f"✅ Заполнил '{selector}' значением: {value[:100]}",
        )
    except RuntimeError as e:
        return ToolResult("browser_fill", False, "", error=str(e))
    except Exception as e:
        return ToolResult("browser_fill", False, "",
                          error=f"Ошибка заполнения: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# RESEARCH TOOLS (handlers) — Internet Reasoning
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_research(
    query: str,
    max_sources: int = 5,
    **kwargs,
) -> ToolResult:
    """
    Исследовать вопрос с проверкой множества источников.
    Использует Internet Reasoning Engine: поиск, анализ,
    извлечение фактов, обнаружение противоречий, синтез ответа.
    """
    from pds_ultimate.core.internet_reasoning import reasoning_engine

    try:
        answer = await reasoning_engine.research(
            query=query,
            max_sources=int(max_sources),
            expand_queries=True,
        )

        lines = [answer.summary]
        lines.append(f"\n📊 Уверенность: {answer.confidence:.0%}")
        lines.append(f"📖 Источников: {answer.sources_count}")
        lines.append(f"🏷️ Качество: {answer.quality_label}")

        if answer.has_contradictions:
            lines.append(
                f"⚠️ Противоречий: {len(answer.contradictions)}"
            )

        return ToolResult(
            "research", True, "\n".join(lines),
            data=answer.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "research", False, "",
            error=f"Ошибка исследования: {e}",
        )


async def tool_deep_research(
    query: str,
    max_sources: int = 10,
    **kwargs,
) -> ToolResult:
    """
    Глубокое исследование с расширенным покрытием источников.
    Для сложных вопросов, где нужна проверка из множества
    независимых источников.
    """
    from pds_ultimate.core.internet_reasoning import reasoning_engine

    try:
        answer = await reasoning_engine.deep_research(
            query=query,
            max_sources=int(max_sources),
        )

        lines = [answer.summary]
        lines.append(f"\n📊 Уверенность: {answer.confidence:.0%}")
        lines.append(f"📖 Источников: {answer.sources_count}")
        lines.append(f"🔬 Фактов проанализировано: {len(answer.facts)}")
        lines.append(f"🏷️ Качество: {answer.quality_label}")

        if answer.has_contradictions:
            lines.append(
                f"⚠️ Противоречий: {len(answer.contradictions)}"
            )

        stats = reasoning_engine.get_stats()
        lines.append(
            f"\n📈 Статистика: {stats['queries']} запросов, "
            f"{stats['pages']} стр, {stats['time_ms']}мс"
        )

        return ToolResult(
            "deep_research", True, "\n".join(lines),
            data=answer.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "deep_research", False, "",
            error=f"Ошибка глубокого исследования: {e}",
        )


async def tool_quick_search(
    query: str,
    **kwargs,
) -> ToolResult:
    """
    Быстрый поиск без расширения запросов.
    Для простых вопросов, когда нужен быстрый ответ.
    """
    from pds_ultimate.core.internet_reasoning import reasoning_engine

    try:
        answer = await reasoning_engine.quick_search(query=query)

        lines = [answer.summary]
        lines.append(f"\n📊 Уверенность: {answer.confidence:.0%}")
        lines.append(f"📖 Источников: {answer.sources_count}")

        return ToolResult(
            "quick_search", True, "\n".join(lines),
            data=answer.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "quick_search", False, "",
            error=f"Ошибка быстрого поиска: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 10: KNOWLEDGE BASE / SEMANTIC SEARCH V2 (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_knowledge_add(
    content: str,
    category: str = "general",
    source: str = "",
    tags: str = "",
    **kwargs,
) -> ToolResult:
    """Добавить знание в базу знаний."""
    from pds_ultimate.core.semantic_search_v2 import semantic_search_v2

    try:
        tag_list = [t.strip()
                    for t in tags.split(",") if t.strip()] if tags else []
        item_id = semantic_search_v2.add_knowledge(
            content=content,
            category=category,
            source=source,
            tags=tag_list,
        )
        return ToolResult(
            "knowledge_add", True,
            f"📚 Знание добавлено в базу!\n"
            f"  🆔 ID: {item_id}\n"
            f"  📁 Категория: {category}\n"
            f"  🏷️ Теги: {', '.join(tag_list) if tag_list else '—'}",
            data={"id": item_id, "category": category},
        )
    except Exception as e:
        return ToolResult(
            "knowledge_add", False, "",
            error=f"Ошибка добавления знания: {e}",
        )


async def tool_knowledge_search(
    query: str,
    category: str = "",
    max_results: int = 5,
    **kwargs,
) -> ToolResult:
    """Семантический поиск по базе знаний."""
    from pds_ultimate.core.semantic_search_v2 import semantic_search_v2

    try:
        results = semantic_search_v2.search_knowledge(
            query=query,
            category=category or None,
            max_results=int(max_results),
        )
        if not results:
            return ToolResult(
                "knowledge_search", True,
                "🔍 Ничего не найдено в базе знаний.",
                data={"results": [], "count": 0},
            )

        lines = [f"🔍 Найдено {len(results)} результатов:"]
        for i, r in enumerate(results, 1):
            lines.append(
                f"\n  {i}. [{r.item.category.value}] "
                f"(скор: {r.final_score:.2f})\n"
                f"     {r.item.content[:150]}..."
            )
        stats = semantic_search_v2.get_stats()
        lines.append(
            f"\n📊 Всего в базе: {stats['knowledge_base']['total']} знаний")

        return ToolResult(
            "knowledge_search", True, "\n".join(lines),
            data={"results": [r.to_dict() for r in results],
                  "count": len(results)},
        )
    except Exception as e:
        return ToolResult(
            "knowledge_search", False, "",
            error=f"Ошибка поиска: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 10: CONFIDENCE TRACKER (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_confidence_check(
    text: str,
    source_count: int = 1,
    source_agreement: float = 0.5,
    **kwargs,
) -> ToolResult:
    """Оценить уверенность в ответе."""
    from pds_ultimate.core.confidence_tracker import confidence_tracker

    try:
        score = confidence_tracker.estimate(
            text=text,
            source_count=int(source_count),
            source_agreement=float(source_agreement),
        )
        needs = confidence_tracker.needs_search(score)

        lines = [
            f"{score.emoji} Уверенность: {score.value:.0%} ({score.level.value})",
        ]
        if score.factors:
            lines.append("📊 Факторы:")
            for k, v in score.factors.items():
                lines.append(f"  • {k}: {v:.2f}")
        if score.uncertainties:
            lines.append("⚠️ Неопределённости:")
            for u in score.uncertainties:
                lines.append(f"  • {u.value}")
        if needs:
            lines.append("🔍 Рекомендуется дополнительный поиск!")
            plan = confidence_tracker.get_search_plan(score)
            if plan:
                lines.append(f"  План: {plan.get('action', '?')}")

        return ToolResult(
            "confidence_check", True, "\n".join(lines),
            data=score.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "confidence_check", False, "",
            error=f"Ошибка оценки уверенности: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 10: ADAPTIVE QUERY EXPANSION (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_expand_query(
    query: str,
    context: str = "",
    strategy: str = "synonym",
    **kwargs,
) -> ToolResult:
    """Расширить поисковый запрос."""
    from pds_ultimate.core.adaptive_query import adaptive_query

    try:
        expanded = adaptive_query.expand(
            query=query,
            context=context,
            strategy=strategy,
        )
        lines = [
            "🔄 Расширение запроса:",
            f"  📝 Оригинал: {expanded.original}",
            f"  ✨ Расширенный: {expanded.expanded}",
            f"  📋 Стратегия: {expanded.strategy.value}",
            f"  📊 Уверенность: {expanded.confidence:.0%}",
        ]
        if expanded.added_terms:
            lines.append(f"  ➕ Добавлено: {', '.join(expanded.added_terms)}")
        if expanded.removed_terms:
            lines.append(f"  ➖ Убрано: {', '.join(expanded.removed_terms)}")

        return ToolResult(
            "expand_query", True, "\n".join(lines),
            data=expanded.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "expand_query", False, "",
            error=f"Ошибка расширения запроса: {e}",
        )


async def tool_find_gaps(
    query: str,
    answer: str,
    confidence: float = 0.5,
    **kwargs,
) -> ToolResult:
    """Найти пробелы в ответе."""
    from pds_ultimate.core.adaptive_query import adaptive_query

    try:
        gaps = adaptive_query.find_gaps(
            query=query,
            answer=answer,
            confidence=float(confidence),
        )
        if not gaps:
            return ToolResult(
                "find_gaps", True,
                "✅ Пробелов не найдено — ответ полный!",
                data={"gaps": [], "count": 0},
            )

        lines = [f"🔍 Найдено {len(gaps)} пробелов:"]
        for i, gap in enumerate(gaps, 1):
            lines.append(
                f"\n  {i}. [{gap.gap_type.value}] {gap.description}\n"
                f"     Приоритет: {gap.priority:.0%}"
            )
            if gap.suggested_query:
                lines.append(f"     💡 Запрос: {gap.suggested_query}")

        return ToolResult(
            "find_gaps", True, "\n".join(lines),
            data={"gaps": [g.to_dict() for g in gaps], "count": len(gaps)},
        )
    except Exception as e:
        return ToolResult(
            "find_gaps", False, "",
            error=f"Ошибка анализа пробелов: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 10: TASK PRIORITIZER (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_task_add(
    name: str,
    priority: str = "medium",
    task_type: str = "general",
    deadline_sec: float = 0,
    **kwargs,
) -> ToolResult:
    """Добавить задачу в очередь."""
    from pds_ultimate.core.task_prioritizer import task_prioritizer

    try:
        dl = float(deadline_sec) if float(deadline_sec) > 0 else None
        task = task_prioritizer.add_task(
            name=name,
            priority=priority,
            task_type=task_type,
            deadline_sec=dl,
        )
        lines = [
            "📋 Задача добавлена в очередь!",
            f"  🆔 ID: {task.id}",
            f"  📌 Приоритет: {task.priority.name}",
            f"  📁 Тип: {task.task_type}",
        ]
        if task.deadline:
            ttd = task.time_to_deadline
            if ttd is not None:
                lines.append(f"  ⏰ Дедлайн через: {ttd:.0f} сек")
        stats = task_prioritizer.get_stats()
        lines.append(
            f"\n📊 В очереди: {stats['queue']['pending']} задач"
        )
        return ToolResult(
            "task_add", True, "\n".join(lines),
            data=task.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "task_add", False, "",
            error=f"Ошибка добавления задачи: {e}",
        )


async def tool_task_queue(
    action: str = "list",
    **kwargs,
) -> ToolResult:
    """Показать очередь задач."""
    from pds_ultimate.core.task_prioritizer import task_prioritizer

    try:
        if action == "next":
            task = task_prioritizer.next_task()
            if task is None:
                return ToolResult(
                    "task_queue", True,
                    "📋 Очередь пуста — нет задач.",
                    data={"task": None},
                )
            return ToolResult(
                "task_queue", True,
                f"▶️ Следующая задача: {task.name}\n"
                f"  🆔 {task.id} | 📌 {task.priority.name}",
                data=task.to_dict(),
            )

        if action == "plan":
            plan = task_prioritizer.get_plan()
            if not plan:
                return ToolResult(
                    "task_queue", True,
                    "📋 Нет задач для планирования.",
                    data={"plan": []},
                )
            lines = ["📋 План выполнения:"]
            for i, wave in enumerate(plan, 1):
                lines.append(f"\n  🌊 Волна {i} ({len(wave)} задач):")
                for t in wave:
                    lines.append(f"    • {t['name']} [{t['priority']}]")
            est = task_prioritizer.estimate_time()
            lines.append(f"\n⏱️ Оценка времени: {est:.1f} сек")
            return ToolResult(
                "task_queue", True, "\n".join(lines),
                data={"plan": plan, "estimated_sec": est},
            )

        if action == "stats":
            stats = task_prioritizer.get_stats()
            q = stats["queue"]
            lines = [
                "📊 Статистика очереди:",
                f"  📋 Всего: {q['total']}",
                f"  ⏳ Ожидают: {q['pending']}",
                f"  ▶️ Выполняются: {q['running']}",
                f"  ✅ Завершены: {q['completed']}",
                f"  ❌ Ошибки: {q['failed']}",
                f"  ⚠️ Просрочены: {q['overdue']}",
            ]
            return ToolResult(
                "task_queue", True, "\n".join(lines),
                data=stats,
            )

        # Default: list
        stats = task_prioritizer.get_stats()
        q = stats["queue"]
        return ToolResult(
            "task_queue", True,
            f"📋 Очередь задач: {q['pending']} ожидают, "
            f"{q['running']} выполняются, {q['completed']} завершены",
            data=stats,
        )
    except Exception as e:
        return ToolResult(
            "task_queue", False, "",
            error=f"Ошибка очереди задач: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 10: CONTEXT COMPRESSOR (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_summarize_text(
    text: str,
    ratio: float = 0.3,
    recursive: bool = False,
    **kwargs,
) -> ToolResult:
    """Суммаризировать текст."""
    from pds_ultimate.core.context_compressor import context_compressor

    try:
        ratio_val = max(0.1, min(0.9, float(ratio)))
        if recursive or len(text) > 3000:
            result = context_compressor.summarize_recursive(text)
        else:
            result = context_compressor.summarize(text, ratio=ratio_val)

        lines = [
            "📝 Суммаризация:",
            f"  📏 Оригинал: {result.original_length} символов",
            f"  📐 Сжато: {result.compressed_length} символов",
            f"  💾 Экономия: {result.savings_pct:.1f}%",
            f"  📋 Метод: {result.method}",
        ]
        if result.key_terms:
            lines.append(f"  🏷️ Ключевые: {', '.join(result.key_terms[:5])}")
        lines.append(f"\n{result.text}")

        return ToolResult(
            "summarize_text", True, "\n".join(lines),
            data=result.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "summarize_text", False, "",
            error=f"Ошибка суммаризации: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART 10: TIME & RELEVANCE (handlers)
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_check_freshness(
    text: str,
    **kwargs,
) -> ToolResult:
    """Проверить актуальность данных."""
    from pds_ultimate.core.time_relevance import time_relevance

    try:
        report = time_relevance.check_freshness(text)

        lines = [
            f"{report.grade.emoji} Свежесть: {report.grade.value.upper()}",
            f"  📊 Скор: {report.score:.0%}",
            f"  📅 Возраст: {report.data_age_days:.0f} дней",
        ]
        if report.markers:
            lines.append(f"  🔍 Дат найдено: {len(report.markers)}")
            for m in report.markers[:3]:
                lines.append(f"    • «{m.text}» → {m.scope.value}")
        if report.recommendation:
            lines.append(f"\n💡 {report.recommendation}")
        if report.needs_update:
            lines.append("⚠️ Рекомендуется обновить данные!")

        return ToolResult(
            "check_freshness", True, "\n".join(lines),
            data=report.to_dict(),
        )
    except Exception as e:
        return ToolResult(
            "check_freshness", False, "",
            error=f"Ошибка проверки свежести: {e}",
        )


async def tool_time_decay(
    score: float,
    age_days: float,
    method: str = "exponential",
    **kwargs,
) -> ToolResult:
    """Применить временное затухание."""
    from pds_ultimate.core.time_relevance import time_relevance

    try:
        adjusted = time_relevance.apply_time_decay(
            score=float(score),
            age_days=float(age_days),
            method=method,
        )
        delta = adjusted - float(score)
        lines = [
            "⏱️ Временное затухание:",
            f"  📊 Исходный скор: {float(score):.3f}",
            f"  📅 Возраст: {float(age_days):.0f} дней",
            f"  📈 Метод: {method}",
            f"  🎯 Скорректированный: {adjusted:.3f}",
            f"  📉 Дельта: {delta:+.3f}",
        ]
        return ToolResult(
            "time_decay", True, "\n".join(lines),
            data={
                "original": float(score),
                "adjusted": round(adjusted, 4),
                "delta": round(delta, 4),
                "method": method,
                "age_days": float(age_days),
            },
        )
    except Exception as e:
        return ToolResult(
            "time_decay", False, "",
            error=f"Ошибка затухания: {e}",
        )


# ── Part 11: Integration Layer handlers ──────────────────────────────

async def tool_run_chain(
    chain_name: str,
    query: str = "",
    **kwargs,
) -> ToolResult:
    """Запустить цепочку инструментов."""
    from pds_ultimate.core.integration_layer import integration_layer

    try:
        result = await integration_layer.execute_chain(
            chain_name, {"query": query} if query else {},
        )
        if result is None:
            return ToolResult(
                "run_chain", False, "",
                error=f"Цепочка '{chain_name}' не найдена. "
                "Используйте list_chains для списка.",
            )
        lines = [
            f"🔗 Цепочка: {chain_name}",
            f"  📊 Статус: {result.status.value}",
            f"  ⏱️ Время: {result.total_time:.2f}с",
            f"  📋 Шагов: {len(result.step_results)}",
        ]
        for i, sr in enumerate(result.step_results, 1):
            icon = "✅" if sr.success else "❌"
            lines.append(f"  {icon} Шаг {i}: {sr.step_name} "
                         f"({sr.duration:.2f}с)")
        return ToolResult(
            "run_chain", result.success, "\n".join(lines),
            data={
                "chain": chain_name,
                "status": result.status.value,
                "success": result.success,
                "total_time": round(result.total_time, 3),
                "steps": len(result.step_results),
            },
        )
    except Exception as e:
        return ToolResult(
            "run_chain", False, "",
            error=f"Ошибка выполнения цепочки: {e}",
        )


async def tool_health_check(
    action: str = "report",
    **kwargs,
) -> ToolResult:
    """Показать здоровье инструментов."""
    from pds_ultimate.core.integration_layer import integration_layer

    try:
        if action == "stats":
            stats = integration_layer.get_stats()
            lines = [
                "📊 Статистика интеграции:",
                f"  🔗 Цепочек: {stats.get('chains', 0)}",
                f"  🛡️ Breakers: {stats.get('circuit_breakers', 0)}",
                f"  📈 Метрик: {stats.get('metrics', 0)}",
                f"  🔄 Fallbacks: {stats.get('fallbacks', 0)}",
                f"  🩺 Auto-heals: {stats.get('auto_heals', 0)}",
            ]
            return ToolResult(
                "tool_health", True, "\n".join(lines), data=stats,
            )

        report = integration_layer.get_health_report()
        if action == "unhealthy":
            report = {k: v for k, v in report.items()
                      if v.get("health") != "healthy"}
        elif action == "slow":
            report = {k: v for k, v in report.items()
                      if v.get("avg_time", 0) > 2.0}

        if not report:
            return ToolResult(
                "tool_health", True,
                "✅ Все инструменты работают нормально.",
                data={"healthy": True},
            )

        lines = [f"🩺 Здоровье инструментов ({len(report)}):"]
        for name, info in list(report.items())[:20]:
            health = info.get("health", "unknown")
            icon = {"healthy": "✅", "degraded": "⚠️",
                    "unhealthy": "❌"}.get(health, "❓")
            lines.append(f"  {icon} {name}: {health}")
        return ToolResult(
            "tool_health", True, "\n".join(lines), data=report,
        )
    except Exception as e:
        return ToolResult(
            "tool_health", False, "",
            error=f"Ошибка проверки здоровья: {e}",
        )


async def tool_parallel_execute(
    calls: str,
    **kwargs,
) -> ToolResult:
    """Выполнить несколько инструментов параллельно."""
    from pds_ultimate.core.integration_layer import integration_layer

    try:
        # Парсим формат: tool1:p1=v1,p2=v2;tool2:p1=v1
        parsed = []
        for part in calls.split(";"):
            part = part.strip()
            if not part:
                continue
            if ":" in part:
                tname, params_str = part.split(":", 1)
                params = {}
                for kv in params_str.split(","):
                    kv = kv.strip()
                    if "=" in kv:
                        k, v = kv.split("=", 1)
                        params[k.strip()] = v.strip()
                parsed.append((tname.strip(), params))
            else:
                parsed.append((part.strip(), {}))

        if not parsed:
            return ToolResult(
                "parallel_tools", False, "",
                error="Не указаны вызовы. Формат: tool1:p1=v1;tool2:p2=v2",
            )

        results = await integration_layer.execute_parallel(parsed)
        ok = sum(1 for r in results if getattr(r, "success", False))
        lines = [
            f"⚡ Параллельное выполнение: {ok}/{len(results)} успешно",
        ]
        for i, r in enumerate(results):
            tname = parsed[i][0] if i < len(parsed) else "?"
            icon = "✅" if getattr(r, "success", False) else "❌"
            out = getattr(r, "output", "")
            snippet = (out[:60] + "…") if len(out) > 60 else out
            lines.append(f"  {icon} {tname}: {snippet}")
        return ToolResult(
            "parallel_tools", True, "\n".join(lines),
            data={"total": len(results), "success": ok},
        )
    except Exception as e:
        return ToolResult(
            "parallel_tools", False, "",
            error=f"Ошибка параллельного выполнения: {e}",
        )


async def tool_list_chains(**kwargs) -> ToolResult:
    """Показать все доступные цепочки."""
    from pds_ultimate.core.integration_layer import integration_layer

    try:
        chains = list(integration_layer.chains.keys())
        router_chains = list(integration_layer.router.routes.keys()) \
            if integration_layer.router else []
        lines = [f"🔗 Доступные цепочки ({len(chains)}):"]
        for ch in chains:
            chain = integration_layer.chains[ch]
            lines.append(f"  • {ch} ({len(chain.steps)} шагов)")
        if router_chains:
            lines.append(f"\n🗺️ Авто-маршруты ({len(router_chains)}):")
            for rc in router_chains:
                lines.append(f"  • {rc}")
        return ToolResult(
            "list_chains", True, "\n".join(lines),
            data={"chains": chains, "routes": router_chains},
        )
    except Exception as e:
        return ToolResult(
            "list_chains", False, "",
            error=f"Ошибка получения списка цепочек: {e}",
        )


# ── Part 12: Production Hardening handlers ───────────────────────────────

async def tool_system_health(
    section: str = "full",
    **kwargs,
) -> ToolResult:
    """Полный системный отчёт."""
    from pds_ultimate.core.production import production

    try:
        report = production.get_system_report()

        if section != "full" and section in report:
            report = {section: report[section]}

        lines = ["🏥 Системный отчёт:"]

        # Uptime
        if "uptime" in report:
            up = report["uptime"]
            lines.append(f"  ⏱️ Аптайм: {up.get('uptime_human', '?')}")
            lines.append(f"  🔄 Перезагрузок: {up.get('restarts', 0)}")

        # Health
        if "health" in report:
            h = report["health"]
            overall = h.get("overall", "unknown")
            icon = {"healthy": "✅", "degraded": "⚠️",
                    "unhealthy": "❌"}.get(overall, "❓")
            lines.append(f"  {icon} Здоровье: {overall}")
            subs = h.get("subsystems", {})
            for name, info in list(subs.items())[:10]:
                s_icon = {"healthy": "✅", "degraded": "⚠️",
                          "unhealthy": "❌"}.get(
                    info.get("status", ""), "❓")
                lines.append(f"    {s_icon} {name}: {info.get('status', '?')}")

        # System
        if "system" in report:
            sys_m = report["system"]
            mem = sys_m.get("memory", {})
            disk = sys_m.get("disk", {})
            if mem.get("rss_mb"):
                lines.append(f"  💾 RAM: {mem['rss_mb']}MB")
            if disk.get("free_gb"):
                lines.append(
                    f"  💿 Диск: {disk.get('usage_percent', 0)}% "
                    f"({disk['free_gb']}GB свободно)")

        # Requests
        if "requests" in report:
            req = report["requests"]
            lines.append(
                f"  📊 Запросов: {req.get('total_requests', 0)} "
                f"(ошибок: {req.get('error_rate', 0)}%)")

        # Alerts
        if "alerts" in report:
            active = report["alerts"].get("active", [])
            if active:
                lines.append(f"  🚨 Активных алертов: {len(active)}")
                for a in active[:5]:
                    lines.append(f"    ⚠️ {a.get('name', '?')}: "
                                 f"{a.get('message', '')}")
            else:
                lines.append("  ✅ Алертов нет")

        return ToolResult(
            "system_health", True, "\n".join(lines), data=report,
        )
    except Exception as e:
        return ToolResult(
            "system_health", False, "",
            error=f"Ошибка системного отчёта: {e}",
        )


async def tool_rate_limit_info(
    key: str = "",
    **kwargs,
) -> ToolResult:
    """Статус rate-лимитов."""
    from pds_ultimate.core.production import production

    try:
        if key:
            status = production.rate_limiter.get_status(key)
            lines = [
                f"🚦 Rate limit для '{key}':",
                f"  📊 Запросов: {status.get('current_count', 0)}"
                f"/{status.get('max_requests', '?')}",
                f"  ⏳ Осталось: {status.get('remaining', '?')}",
                f"  🚫 Заблокирован: {'да' if status.get('blocked') else 'нет'}",
            ]
        else:
            stats = production.rate_limiter.get_stats()
            lines = [
                "🚦 Rate Limits:",
                f"  📊 Ключей: {stats['total_keys']}",
                f"  🚫 Ограничено: {stats['total_limited']}",
                f"  ⛔ Заблокировано: {stats['currently_blocked']}",
                f"  ⚙️ Custom лимитов: {stats['custom_limits']}",
            ]
            status = stats

        return ToolResult(
            "rate_limit_info", True, "\n".join(lines), data=status,
        )
    except Exception as e:
        return ToolResult(
            "rate_limit_info", False, "",
            error=f"Ошибка rate limit info: {e}",
        )


async def tool_error_report(
    action: str = "recent",
    **kwargs,
) -> ToolResult:
    """Отчёт об ошибках."""
    from pds_ultimate.core.production import production

    try:
        er = production.error_reporter

        if action == "clear":
            er.clear()
            return ToolResult(
                "error_report", True,
                "🗑️ История ошибок очищена.",
                data={"cleared": True},
            )

        if action == "top":
            top = er.get_top_errors(10)
            if not top:
                return ToolResult(
                    "error_report", True,
                    "✅ Ошибок не зафиксировано.", data={"top": []},
                )
            lines = ["📊 Топ ошибок по частоте:"]
            for t in top:
                lines.append(f"  • {t['type']}: {t['count']} раз")
            return ToolResult(
                "error_report", True, "\n".join(lines), data={"top": top},
            )

        if action == "stats":
            stats = er.get_stats()
            lines = [
                "📊 Статистика ошибок:",
                f"  📈 Всего: {stats['total_errors']}",
                f"  🏷️ Типов: {stats['unique_types']}",
                f"  📍 Источников: {stats['unique_sources']}",
            ]
            return ToolResult(
                "error_report", True, "\n".join(lines), data=stats,
            )

        # recent (default)
        recent = er.get_recent(10)
        if not recent:
            return ToolResult(
                "error_report", True,
                "✅ Недавних ошибок нет.", data={"recent": []},
            )
        lines = [f"🔴 Последние ошибки ({len(recent)}):"]
        for r in recent:
            lines.append(
                f"  • [{r['type']}] {r['message'][:80]} "
                f"({r['ago_s']:.0f}с назад)"
            )
        return ToolResult(
            "error_report", True, "\n".join(lines),
            data={"recent": recent},
        )
    except Exception as e:
        return ToolResult(
            "error_report", False, "",
            error=f"Ошибка отчёта об ошибках: {e}",
        )


async def tool_uptime_info(**kwargs) -> ToolResult:
    """Информация об аптайме."""
    from pds_ultimate.core.production import production

    try:
        stats = production.uptime.get_stats()
        lines = [
            "⏱️ Аптайм системы:",
            f"  🕐 Работает: {stats['uptime_human']}",
            f"  📅 Запущена: {stats['started_at']}",
            f"  🔄 Перезагрузок: {stats['restarts']}",
            f"  ⏸️ Простой: {stats['total_downtime_s']}с",
            f"  💓 Последний heartbeat: {stats['last_heartbeat_ago_s']:.1f}с назад",
        ]
        return ToolResult(
            "uptime_info", True, "\n".join(lines), data=stats,
        )
    except Exception as e:
        return ToolResult(
            "uptime_info", False, "",
            error=f"Ошибка аптайм info: {e}",
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PART v6: PERSONA & PROACTIVE TOOLS
# ═══════════════════════════════════════════════════════════════════════════════


async def tool_persona_stats(**kwargs) -> ToolResult:
    """Статистика персоны: сколько пользователей изучено, группы сходства."""
    try:
        from pds_ultimate.core.persona_engine import persona_engine

        stats = persona_engine.get_stats()
        lines = [
            "🧠 Persona Engine:",
            f"  👥 Пользователей изучено: {stats['users']}",
            f"  🔗 Групп сходства: {stats['shared_groups']}",
            f"  📅 Последний retrain: {stats['last_retrain_at']}",
        ]
        return ToolResult("persona_stats", True, "\n".join(lines), data=stats)
    except Exception as e:
        return ToolResult("persona_stats", False, "", error=str(e))


async def tool_persona_retrain(days: int = 3, **kwargs) -> ToolResult:
    """Принудительный retrain персоны из истории чатов."""
    try:
        from pds_ultimate.core.persona_engine import persona_engine

        # Reset interval to force retrain
        persona_engine._last_retrain_at = 0
        result = persona_engine.run_periodic_retrain(days=days)
        if result.get("retrained"):
            return ToolResult(
                "persona_retrain", True,
                f"✅ Retrain завершён: обработано {result['processed']} сообщений.",
                data=result,
            )
        return ToolResult(
            "persona_retrain", False,
            f"⚠️ Retrain не выполнен: {result.get('reason', 'unknown')}",
        )
    except Exception as e:
        return ToolResult("persona_retrain", False, "", error=str(e))


async def tool_persona_style(chat_id: int = 0, **kwargs) -> ToolResult:
    """Показать стиль-гайд для конкретного пользователя."""
    try:
        from pds_ultimate.config import config
        from pds_ultimate.core.persona_engine import persona_engine

        cid = chat_id or config.telegram.owner_id
        guide = persona_engine.get_style_guide(cid)
        if guide:
            return ToolResult("persona_style", True, guide)
        return ToolResult(
            "persona_style", True,
            "📋 Недостаточно данных для стиль-гайда (нужно > 6 сообщений).",
        )
    except Exception as e:
        return ToolResult("persona_style", False, "", error=str(e))


async def tool_proactive_status(**kwargs) -> ToolResult:
    """Статус проактивного движка: задачи, аномалии, фильтры."""
    try:
        from pds_ultimate.core.proactive_engine import proactive_engine

        stats = proactive_engine.get_stats()
        lines = [
            "⚡ Proactive Engine:",
            f"  ✅ Запущен: {stats.get('running', False)}",
            f"  📋 Задач в очереди: {stats.get('pending_tasks', 0)}",
            f"  📊 Важных ключевых слов: {stats.get('important_keywords', 0)}",
            f"  🔔 Событий за сессию: {stats.get('events_logged', 0)}",
        ]
        return ToolResult("proactive_status", True, "\n".join(lines), data=stats)
    except Exception as e:
        return ToolResult("proactive_status", False, "", error=str(e))


async def tool_add_important_keyword(keyword: str, weight: float = 1.0, **kwargs) -> ToolResult:
    """Добавить ключевое слово для проактивного мониторинга сообщений."""
    try:
        from pds_ultimate.core.proactive_engine import proactive_engine

        proactive_engine.add_important_keyword(keyword)
        return ToolResult(
            "add_important_keyword", True,
            f"✅ Ключевое слово '{keyword}' добавлено (вес: {weight}).",
        )
    except Exception as e:
        return ToolResult("add_important_keyword", False, "", error=str(e))
